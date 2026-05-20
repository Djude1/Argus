from pathlib import Path

from django.conf import settings
from docx import Document

from apps.scans.models import ScanJob


def build_scan_report(scan_job: ScanJob) -> str:
    report_dir = Path(settings.MEDIA_ROOT) / "reports"
    report_dir.mkdir(parents=True, exist_ok=True)
    output_path = report_dir / f"scan-{scan_job.id}-report.docx"

    document = Document()
    document.add_heading("Argus 網站健檢報告", 0)
    document.add_paragraph(f"掃描網址：{scan_job.normalized_url}")
    document.add_paragraph(f"掃描狀態：{scan_job.get_status_display()}")
    overall_score = scan_job.overall_score if scan_job.overall_score is not None else "尚未產生"
    document.add_paragraph(f"整體分數：{overall_score}")

    document.add_heading("摘要", level=1)
    for category, score in (scan_job.category_scores or {}).items():
        document.add_paragraph(f"{category.upper()}：{score}")

    document.add_heading("優先處理項目", level=1)
    for action in scan_job.top_actions or []:
        document.add_paragraph(
            f"{action.get('severity', '').upper()} / {action.get('category', '').upper()}："
            f"{action.get('title', '')}",
            style="List Bullet",
        )

    document.add_heading("Findings", level=1)
    for finding in scan_job.findings.select_related("page").all():
        document.add_heading(finding.title, level=2)
        document.add_paragraph(f"分類：{finding.category} / 嚴重度：{finding.severity}")
        if finding.page:
            document.add_paragraph(f"頁面：{finding.page.final_url}")
        document.add_paragraph("問題描述")
        document.add_paragraph(finding.description)
        document.add_paragraph("修補方向")
        document.add_paragraph(finding.remediation)
        if finding.evidence:
            document.add_paragraph("證據")
            document.add_paragraph(finding.evidence[:1000])

    document.add_heading("附錄", level=1)
    document.add_paragraph("本報告僅提供問題描述、證據與修補方向，不產生修復後程式碼。")
    document.save(output_path)
    return str(output_path)
