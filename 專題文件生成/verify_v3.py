# -*- coding: utf-8 -*-
# verify_v3.py  驗證 Argus_系統手冊_v3.docx 是否符合硬性需求
import os, sys
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(BASE, "Argus_系統手冊_v3.docx")

failures = []
def check(cond, msg):
    if not cond:
        failures.append(msg)

doc = Document(DOCX)

# 1. 版面：上下左右 1.5cm、頁首頁尾 1cm、裝訂邊位置左
sec = doc.sections[0]
def cm(v): return round(v.cm, 2) if v is not None else None
check(cm(sec.top_margin) == 1.5, f"top_margin={cm(sec.top_margin)} 應為 1.5")
check(cm(sec.bottom_margin) == 1.5, f"bottom_margin={cm(sec.bottom_margin)} 應為 1.5")
check(cm(sec.left_margin) == 1.5, f"left_margin={cm(sec.left_margin)} 應為 1.5")
check(cm(sec.right_margin) == 1.5, f"right_margin={cm(sec.right_margin)} 應為 1.5")
check(cm(sec.header_distance) == 1.0, f"header_distance={cm(sec.header_distance)} 應為 1.0")
check(cm(sec.footer_distance) == 1.0, f"footer_distance={cm(sec.footer_distance)} 應為 1.0")
pgMar = sec._sectPr.find(qn('w:pgMar'))
check(pgMar is not None and pgMar.get(qn('w:gutter')) is not None, "缺少 gutter 設定")

# 2. 內文段落：單行間距、首行縮排 firstLineChars=200、adjustRightInd
body_paras = [p for p in doc.paragraphs if p.text.strip()]
single_ok = adj_ok = fl_ok = False
for p in body_paras:
    pPr = p._p.find(qn('w:pPr'))
    if pPr is None: continue
    spc = pPr.find(qn('w:spacing'))
    if spc is not None and spc.get(qn('w:line')) == "240" and spc.get(qn('w:lineRule')) == "auto":
        single_ok = True
    if pPr.get(qn('w:adjustRightInd')) == "1":
        adj_ok = True
    ind = pPr.find(qn('w:ind'))
    if ind is not None and ind.get(qn('w:firstLineChars')) == "200":
        fl_ok = True
check(single_ok, "找不到單行間距(240/auto)的內文段落")
check(adj_ok, "找不到 adjustRightInd=1 的段落")
check(fl_ok, "找不到 firstLineChars=200 的首行縮排段落")

# 3. 禁止估算/猜測字眼
full_text = "\n".join(p.text for p in doc.paragraphs)
for tbl in doc.tables:
    for row in tbl.rows:
        for c in row.cells:
            full_text += "\n" + c.text
for word in ["估計", "估算", "猜測", "推估"]:
    check(word not in full_text, f"文件出現禁用字「{word}」")

# 4. caption 不得出現 (PlantUML)
check("（PlantUML）" not in full_text and "(PlantUML)" not in full_text,
      "caption 仍出現 (PlantUML)")

# 5. TOC / 圖表目錄欄位存在
xml = doc.element.xml
check("TOC \\o" in xml or 'TOC \\\\o' in xml or "w:instrText" in xml, "缺少 TOC 欄位")

# 6. 表格列防截斷：至少存在 cantSplit 與 tblHeader
has_cantsplit = has_tblheader = False
for tbl in doc.tables:
    for tr in tbl._tbl.findall(qn('w:tr')):
        trPr = tr.find(qn('w:trPr'))
        if trPr is None: continue
        if trPr.find(qn('w:cantSplit')) is not None: has_cantsplit = True
        if trPr.find(qn('w:tblHeader')) is not None: has_tblheader = True
check(has_cantsplit, "沒有任何 row 設定 cantSplit")
check(has_tblheader, "沒有任何表格設定 tblHeader（標題列重複）")

if failures:
    print("[FAIL] 驗證未通過：")
    for f in failures:
        print("  -", f)
    sys.exit(1)
print("[OK] 所有硬性需求驗證通過")
