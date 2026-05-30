# -*- coding: utf-8 -*-
# gen_argus_v3.py  生成 Argus 系統手冊 v3（初審版，大學部，第1-8章）
# 執行: uv run --project C:\Users\ntub\Desktop\Argus python gen_argus_v3.py

import os, io, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib import font_manager as fm

# matplotlib 載入標楷體，解決中文亂碼
_KAIU = r"C:\Windows\Fonts\kaiu.ttf"
if os.path.exists(_KAIU):
    fm.fontManager.addfont(_KAIU)
    plt.rcParams["font.family"] = "DFKai-SB"
plt.rcParams["axes.unicode_minus"] = False
_FP = fm.FontProperties(fname=_KAIU) if os.path.exists(_KAIU) else None

# ── 常數 ──────────────────────────────────────────
CH_FONT        = "標楷體"
EN_FONT        = "Times New Roman"
BODY_SIZE      = Pt(14)
SEC_SIZE       = Pt(16)
CHAP_SIZE      = Pt(18)
GANTT_EXPECTED = "D9D9D9"   # 淺灰：預期進度
GANTT_ACTUAL   = "595959"   # 深灰：實際進度
HDR_FILL       = "E4DFEC"   # 淺紫：表格標題列（仿黃金屋淺色系）
ALT_FILL       = "F2F2F2"   # 淺灰：交替列

BASE_DIR         = os.path.dirname(os.path.abspath(__file__))
PLACEHOLDER_IMG  = os.path.join(BASE_DIR, "留空用照片.png")
OUT_DOCX         = os.path.join(BASE_DIR, "Argus_系統手冊_v3.docx")
OUT_PLANTUML     = os.path.join(BASE_DIR, "plantuml_diagrams_v3.txt")

# ── 真實可引用資料（來源見 data_sources_v3.md）─────────
# 競品月費（換算新台幣）：各官方定價頁 2026/05 查得；外幣依台灣銀行牌告匯率
# 約 1 USD≈NT$32、1 EUR≈NT$35 換算（匯率每日浮動，僅供比較）
COST_DATA = [
    ("Argus\n（本系統按需點數）",  300),   # 本系統定價：單次掃描最低約 NT$300
    ("Google\nSearch Console",     0),     # 官方免費
    ("Screaming Frog\n（年費折月）", 714),  # €245/年 ÷12 ×35
    ("Ahrefs\n（Starter 方案）",   928),    # US$29/月 ×32
    ("SEMrush\n（Pro 方案）",     4478),    # US$139.95/月 ×32
]

# ── 基礎排版工具 ────────────────────────────────────

def _set_run_font(run, size=None, bold=False, italic=False, color=None):
    if size is None:
        size = BODY_SIZE
    run.font.name  = EN_FONT
    run.font.size  = size
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), CH_FONT)
    rFonts.set(qn("w:ascii"),    EN_FONT)
    rFonts.set(qn("w:hAnsi"),    EN_FONT)
    old = rPr.find(qn("w:rFonts"))
    if old is not None:
        rPr.remove(old)
    rPr.insert(0, rFonts)


def _set_para_spacing(p, before=0, after=6, ls=None):
    pf = p.paragraph_format
    pf.space_before        = Pt(before)
    pf.space_after         = Pt(after)
    pf.line_spacing_rule   = WD_LINE_SPACING.SINGLE   # 單行間距
    # adjustRightInd=1（文件格線時自動調整右側縮排）
    pPr = p._p.get_or_add_pPr()
    pPr.set(qn('w:adjustRightInd'), "1")
    pPr.set(qn('w:snapToGrid'), "0")


def _cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(shd)


def _cell_write(cell, text, size=Pt(12), bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    _set_run_font(run, size, bold=bold)
    _set_para_spacing(p, before=2, after=2, ls=1.2)


def _cell_add_para(cell, text, size=Pt(11), bold=False,
                   left_indent_cm=0, bullet=""):
    p = cell.add_paragraph()
    if left_indent_cm:
        p.paragraph_format.left_indent = Cm(left_indent_cm)
    run = p.add_run(bullet + text)
    _set_run_font(run, size, bold=bold)
    _set_para_spacing(p, before=0, after=2, ls=1.2)
    return p

# ── 段落層級工具 ────────────────────────────────────

def add_chapter(doc, text):
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_run_font(run, CHAP_SIZE, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)   # 取消 Heading 預設藍色
    _set_para_spacing(p, before=18, after=10)


def add_section(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_run_font(run, SEC_SIZE, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    _set_para_spacing(p, before=12, after=6)


def add_body(doc, text, indent=False, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    _set_run_font(run, BODY_SIZE, bold=bold)
    _set_para_spacing(p, before=0, after=4)
    if indent:
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind'); pPr.append(ind)
        ind.set(qn('w:firstLineChars'), "200")   # 位移 2 字元
    return p


def add_bullet(doc, text, level=0):
    """多層次清單 level=0: ➤  level=1: •"""
    p = doc.add_paragraph()
    indent_cm = 1.0 + level * 0.8
    p.paragraph_format.left_indent      = Cm(indent_cm)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    bullet = "➤ " if level == 0 else "• "
    run = p.add_run(bullet + text)
    _set_run_font(run, BODY_SIZE)
    _set_para_spacing(p, before=0, after=3, ls=1.5)


def _add_field(paragraph, instr, default_text="（請按 F9 更新欄位）"):
    """插入 Word 欄位（begin/instrText/separate/text/end）"""
    run = paragraph.add_run()
    r = run._r
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin'); r.append(fc1)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr; r.append(it)
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'separate'); r.append(fc2)
    t = OxmlElement('w:t'); t.text = default_text; r.append(t)
    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'end'); r.append(fc3)
    _set_run_font(run, BODY_SIZE)


def _add_seq_caption(doc, prefix, chapter_no, text, above):
    """圖/表標號：prefix='圖'|'表'；章碼固定、章內流水以 SEQ 自動編號。
    呈現為『圖 3-1　名稱』（章-SEQ）；套 Caption 樣式供圖/表目錄欄位收集。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.style = doc.styles['Caption']
    except KeyError:
        pass
    run = p.add_run(f"{prefix} {chapter_no}-")
    _set_run_font(run, BODY_SIZE)
    seqrun = p.add_run()
    # SEQ 識別碼固定為「圖」/「表」（供圖表目錄 \c 收集）；\s 1 使每遇 Heading 1 重置 → 章內流水
    for tag, txt in [('begin', None), (None, f' SEQ {prefix} \\s 1 \\* ARABIC '), ('end', None)]:
        if tag:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), tag); seqrun._r.append(fc)
        else:
            it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = txt; seqrun._r.append(it)
    _set_run_font(seqrun, BODY_SIZE)
    tail = p.add_run("　" + text)
    _set_run_font(tail, BODY_SIZE)
    _set_para_spacing(p, before=(12 if above else 4), after=(4 if above else 12))
    return p


def _parse_caption(full):
    """從『圖 3-1-1　名稱』解析出 (prefix, 章碼, 名稱)。"""
    m = re.match(r'^(圖|表)\s*(\d+)-\d+-\d+\s*[　 ]\s*(.*)$', full)
    if m:
        return m.group(1), m.group(2), m.group(3)
    m2 = re.match(r'^(圖|表)\s*[　 ]\s*(.*)$', full)
    return (m2.group(1), "0", m2.group(2)) if m2 else ("圖", "0", full)


def add_fig_caption(doc, text):
    prefix, ch, body = _parse_caption(text)
    _add_seq_caption(doc, prefix, ch, body, above=False)


def add_table_caption(doc, text):
    prefix, ch, body = _parse_caption(text)
    _add_seq_caption(doc, prefix, ch, body, above=True)


def add_placeholder(doc, caption, width=Cm(14)):
    """插入留空圖片佔位符"""
    if os.path.exists(PLACEHOLDER_IMG):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(PLACEHOLDER_IMG, width=width)
        _set_para_spacing(p, before=6, after=0, ls=1.0)
    else:
        add_body(doc, f"【圖片佔位符：{caption}】")
    add_fig_caption(doc, caption)


def _table_no_split(tbl, repeat_header=True, header_rows=1):
    """所有資料列不跨頁斷裂；前 header_rows 列設為跨頁重複標題"""
    rows = tbl._tbl.findall(qn('w:tr'))
    for i, tr in enumerate(rows):
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr'); tr.insert(0, trPr)
        cs = OxmlElement('w:cantSplit'); trPr.append(cs)
        if i < header_rows and repeat_header:
            th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trPr.append(th)


def _table_fixed_widths(tbl, widths_cm):
    """固定欄寬，避免同表欄寬不一致"""
    tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)
    for row in tbl.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = Cm(widths_cm[idx])


def _keep_table_together(tbl):
    """強制整表不跨頁（BMC/SWOT 用）：所有列 cantSplit 且不重複標題"""
    for tr in tbl._tbl.findall(qn('w:tr')):
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr'); tr.insert(0, trPr)
        trPr.append(OxmlElement('w:cantSplit'))


def add_std_table(doc, headers, rows, caption):
    """標準淺色表頭表格，標號在上"""
    add_table_caption(doc, caption)
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 標題列
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_write(c, h, size=Pt(12), bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_shading(c, HDR_FILL)
        for run in c.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0, 0, 0)   # 淺色表頭配黑字
    # 資料列
    for ri, row_data in enumerate(rows):
        row  = tbl.rows[ri + 1]
        fill = ALT_FILL if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _cell_write(c, str(val), size=Pt(12))
            _cell_shading(c, fill)
    _table_no_split(tbl, repeat_header=True)
    doc.add_paragraph()


def set_page_margins(doc):
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(1.5)
        sec.right_margin  = Cm(1.5)
        sec.header_distance = Cm(1.0)
        sec.footer_distance = Cm(1.0)
        # 裝訂邊（gutter）位置左；oxml 直接設定 gutter 屬性，gutterAtTop 不設預設靠左
        pgMar = sec._sectPr.find(qn('w:pgMar'))
        if pgMar is not None:
            pgMar.set(qn('w:gutter'), "0")


def add_page_number(doc):
    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        for tag, text in [("begin", None), (None, " PAGE "), ("end", None)]:
            if tag:
                fc = OxmlElement("w:fldChar")
                fc.set(qn("w:fldCharType"), tag)
                run._r.append(fc)
            else:
                instr = OxmlElement("w:instrText")
                instr.set(qn("xml:space"), "preserve")
                instr.text = text
                run._r.append(instr)
        _set_run_font(run, Pt(12))

# ── 圖表工具 ────────────────────────────────────────

def _make_cost_chart():
    """競品月費比較長條圖（真實官方定價換算）；標號由 caption 提供，圖內不重複"""
    tools = [t for t, _ in COST_DATA]
    costs = [c for _, c in COST_DATA]
    # 淺色系（仿黃金屋範例）
    colors = ["#8E7CC3", "#A2C4C9", "#B6D7A8", "#F9CB9C", "#EA9999"]
    fig, ax = plt.subplots(figsize=(9, 4))
    bars = ax.bar(tools, costs, color=colors, edgecolor="#888888", linewidth=0.6)
    ax.set_ylabel("每月費用（新台幣 NT$）", fontsize=10, fontproperties=_FP)
    ax.set_title("主要競品月費比較", fontsize=12, fontweight="bold", fontproperties=_FP)
    ax.set_ylim(0, max(costs) * 1.2)
    for bar, cost in zip(bars, costs):
        label = "免費" if cost == 0 else f"NT${cost:,}"
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(costs) * 0.015,
                label, ha="center", va="bottom", fontsize=9, fontweight="bold",
                fontproperties=_FP)
    for lbl in ax.get_xticklabels():
        lbl.set_fontproperties(_FP)
        lbl.set_fontsize(9)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"NT${int(x):,}"))
    for lbl in ax.get_yticklabels():
        lbl.set_fontproperties(_FP)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _insert_chart(doc, buf, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Cm(14))
    _set_para_spacing(p, before=6, after=0, ls=1.0)
    buf.close()
    add_fig_caption(doc, caption)

# ── 封面頁 ──────────────────────────────────────────

def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()

    def ctr(text, size, bold=True):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _set_run_font(run, Pt(size), bold=bold)
        _set_para_spacing(p, before=6, after=6, ls=1.5)

    ctr("國立臺北商業大學", 22)
    ctr("資 訊 管 理 系", 22)
    doc.add_paragraph()
    ctr("115' 資訊系統專案設計", 18)
    doc.add_paragraph()
    ctr("系 統 手 冊", 28)
    doc.add_paragraph()

    info = [
        ("組　　別：", "第 115○○○ 組"),
        ("題　　目：", "Argus 網站健檢與資安掃描 SaaS 平台"),
        ("指導老師：", "○○○ 老師"),
        ("組　　長：", "○○○○○  ○○○"),
        ("組　　員：", "○○○○○  ○○○"),
        ("",           "○○○○○  ○○○"),
        ("",           "○○○○○  ○○○"),
    ]
    for label, val in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if label:
            r1 = p.add_run(label)
            _set_run_font(r1, Pt(16), bold=True)
        r2 = p.add_run(val)
        _set_run_font(r2, Pt(16))
        _set_para_spacing(p, before=2, after=2, ls=1.5)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("中華民國 115 年 ○ 月 ○○ 日")
    _set_run_font(run, Pt(16), bold=True)
    doc.add_page_break()

# ── 目錄 / 圖目錄 / 表目錄 ────────────────────────────

def add_toc_pages(doc):
    def heading(title):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        _set_run_font(run, CHAP_SIZE, bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
        _set_para_spacing(p, before=12, after=12)

    # ── 目錄（TOC 欄位，收集 Heading 1-2）──
    heading("目　　錄")
    p = doc.add_paragraph()
    _add_field(p, 'TOC \\o "1-2" \\h \\z \\u')
    doc.add_page_break()

    # ── 圖目錄（Table of Figures，收集 SEQ 圖）──
    heading("圖 目 錄")
    p = doc.add_paragraph()
    _add_field(p, 'TOC \\h \\z \\c "圖"')
    doc.add_page_break()

    # ── 表目錄（Table of Figures，收集 SEQ 表）──
    heading("表 目 錄")
    p = doc.add_paragraph()
    _add_field(p, 'TOC \\h \\z \\c "表"')
    doc.add_page_break()

# ── 第1章　前言 ─────────────────────────────────────

def ch1(doc):
    add_chapter(doc, "第1章　前言")

    # 1-1 背景介紹
    add_section(doc, "1-1　背景介紹")
    add_body(doc,
        "隨著數位轉型浪潮席捲全球，企業網站已成為品牌形象與商業交易的核心載體。"
        "根據 Statista（2024）統計，全球網站數量已突破 17 億個，Google 每日索引超過 1,300 億個網頁。"
        "在此背景下，網站健康狀況直接影響搜尋引擎排名、使用者體驗與資訊安全防護能力。"
        "傳統的網站健檢工具多以單一維度切入，未能全面涵蓋 SEO（搜尋引擎優化）、"
        "AEO（答案引擎優化）、GEO（生成式引擎優化）與資安掃描等多元需求，"
        "且設置門檻較高，不利於中小企業自行部署。", indent=True)
    add_body(doc,
        "與此同時，生成式 AI（如 ChatGPT、Gemini、Perplexity）的崛起正在重塑資訊搜尋行為。"
        "根據 BrightEdge（2023）報告，約 54% 的美國成年人已開始透過 AI 聊天機器人取代傳統搜尋引擎。"
        "此一趨勢催生了對 AEO 與 GEO 優化的迫切需求——企業不僅需要在 Google 排名，"
        "更需要確保自身內容能被 AI 正確理解並引用。"
        "資安威脅亦持續升溫，根據 OWASP（2021）報告，注入攻擊、認證缺陷、"
        "安全性設定錯誤仍為最常見的網站安全問題，中小企業因缺乏專業資源而特別脆弱。", indent=True)
    add_body(doc,
        "綜合上述背景，市場迫切需要一套能同時解決 SEO、AEO、GEO 與資安檢測的整合型網站健檢平台，"
        "以低廉的成本與友善的操作介面，讓中小企業得以掌握網站的全面健康狀況。", indent=True)

    # 1-2 動機
    add_section(doc, "1-2　動機")
    add_body(doc,
        "本專題的動機源於對現有網站健檢工具碎片化問題的深刻觀察。"
        "市場上的主流工具各有所長但各有局限：SEO 工具（Ahrefs、SEMrush）月費高達數千元新台幣，"
        "資安掃描工具（Nmap、Dirsearch、Katana）操作門檻高且僅提供命令列介面，"
        "使用者需同時訂閱並熟悉多套工具才能獲得完整的診斷視野。"
        "此外，幾乎所有現有工具均未將 AEO 與 GEO 評估納入分析框架，"
        "而這兩個維度在生成式 AI 時代已成為網站競爭力的關鍵指標。", indent=True)
    add_body(doc,
        "有鑑於此，本團隊開發 Argus——命名源自希臘神話中擁有百眼的守衛者，象徵對網站全方位、"
        "無死角的監測能力。Argus 結合 BFS 網頁爬蟲、四維掃描引擎（SEO / AEO / GEO / 資安）"
        "與 AI 代理自動化測試，提供圖形化 Web 介面與按需點數計費模式，"
        "旨在以一站式服務降低網站健檢的技術門檻與使用成本。", indent=True)

    add_body(doc,
        "為使比較基準對等，本表選取與 Argus 同屬「雲端 SaaS 網站分析平台」之主流商業產品"
        "進行功能維度比較，而非單一用途的命令列工具。比較對象涵蓋三大 SEO 分析平台"
        "（SEMrush、Ahrefs、Moz Pro）與一資安監測平台（Sucuri），如表 1-2-1 所示：")

    add_std_table(doc,
        ["功能維度", "Argus（本系統）", "SEMrush", "Ahrefs", "Moz Pro", "Sucuri"],
        [
            ("平台定位",       "四維網站健檢平台", "SEO／數位行銷套件", "SEO／外鏈分析", "SEO 分析平台", "網站資安防護"),
            ("部署形式",       "雲端 SaaS",        "雲端 SaaS",        "雲端 SaaS",   "雲端 SaaS",   "雲端 SaaS"),
            ("SEO 健檢",       "✅ 完整",          "✅ 完整",          "✅ 完整",     "✅ 完整",     "❌"),
            ("AEO / GEO 掃描", "✅ 首創四維評分",  "❌",               "❌",          "❌",          "❌"),
            ("資安被動掃描",   "✅ HTTP 標頭／TLS","❌",               "❌",          "❌",          "✅ 惡意程式／WAF"),
            ("一站式整合四維", "✅ SEO+AEO+GEO+資安","❌（僅 SEO）",   "❌（僅 SEO）", "❌（僅 SEO）","❌（僅資安）"),
            ("圖形化報告匯出", "✅ Word 報告",     "✅",               "✅",          "✅",          "✅ 資安報告"),
            ("中文操作介面",   "✅ 全中文",        "部分",             "部分",        "❌",          "部分"),
            ("計費模式",       "按需點數制",       "月訂閱",           "月訂閱",      "月訂閱",      "年訂閱"),
            ("最低月費（約）", "NT$300 起",        "NT$4,478",         "NT$928",      "NT$1,568",    "NT$533"),
        ],
        "表 1-2-1　主流網站分析 SaaS 平台功能特性比較表")

    add_body(doc,
        "資料來源（各平台官方定價頁，2026 年 5 月查得；外幣依台灣銀行牌告匯率"
        "https://rate.bot.com.tw/xrt 約 1 USD≈NT$32 換算）：SEMrush Pro US$139.95/月"
        "（https://www.semrush.com/pricing/seo/）；Ahrefs Starter US$29/月"
        "（https://ahrefs.com/pricing）；Moz Pro Starter US$49/月"
        "（https://moz.com/products/pro/pricing）；Sucuri Basic Platform US$199.99/年"
        "（https://sucuri.net/website-security-platform/signup/，折合約 NT$533/月）；"
        "Argus 為本系統按需點數定價與系統實測（2026）。",
        indent=True)

    add_body(doc,
        "由比較表可見，Argus 是市場上唯一同時具備四維掃描、圖形化介面、完全被動合規操作"
        "以及按需計費的整合型平台，填補了現有工具在 AEO/GEO 維度的空白。", indent=True)

    # 1-3 系統目的與目標
    add_section(doc, "1-3　系統目的與目標")
    add_body(doc, "本系統的主要目的在於提供中小企業與數位行銷人員一套低門檻、"
        "高整合度的網站健康診斷平台，具體目標如下：", indent=True)
    add_bullet(doc, "建立統一的四維掃描架構：同時涵蓋 SEO 技術指標（Meta 標籤、結構化資料、頁面速度）、"
        "AEO 答案引擎優化（FAQ 結構、精選摘要適配性）、GEO 生成式引擎優化（語意密度、"
        "AI 可讀性評分）及資安被動檢測（HTTPS/HSTS/CSP/XFO 安全標頭），輸出 0–100 分綜合評分報告。")
    add_bullet(doc, "提供按需點數計費的 SaaS 訂閱模式：以 hold→settle→refund 原子交易機制確保"
        "計費冪等安全，支援月贈點與多種購買方案，讓中小企業以 1/10 成本享受企業級診斷服務。")
    add_bullet(doc, "整合 BFS 網頁爬蟲與 Playwright 無頭瀏覽器：以廣度優先搜尋策略深入爬取"
        "網站所有可達頁面（最大深度 max_depth、最大頁數 max_pages），確保掃描覆蓋率完整。")
    add_bullet(doc, "輸出專業 Word 格式健檢報告（.docx）：包含各維度發現事項、建議改善行動"
        "及頁面截圖，可直接提交行銷或工程團隊執行改善。")
    add_bullet(doc, "提供 ReactFlow 拓樸圖視覺化：呈現爬蟲涵蓋的網站頁面結構與連結關係，"
        "協助使用者掌握網站架構全貌。")
    add_bullet(doc, "整合 Hermes-Agent 多模態 AI 代理（Phase 2）：採用 observe→think→act 循環，"
        "透過工具調用（function calling）模擬真實使用者操作，自動偵測表單、按鈕等 UX 瑕疵。")

    # 1-4 預期成果
    add_section(doc, "1-4　預期成果")
    add_body(doc, "本專題預期完成下列具體成果：", indent=True)
    add_bullet(doc, "完成 Django 5.1 + DRF 3.15 後端 API 系統，提供掃描任務管理、點數計費、"
        "使用者認證（Google OAuth 2.0）與管理員後台等完整 RESTful 端點。")
    add_bullet(doc, "完成 React 18 + Vite 5 前端 SPA，涵蓋掃描提交、結果詳情、ReactFlow 拓樸圖、"
        "點數購買與後台管理等完整使用者流程。")
    add_bullet(doc, "完成 Playwright 1.45 BFS 爬蟲引擎，支援最大深度設定與無頭截圖存檔，"
        "並透過 Celery + Redis 實現非同步任務派發。")
    add_bullet(doc, "完成四維掃描引擎（SEO / AEO / GEO / Security），對每一頁面輸出"
        "結構化 Finding 紀錄（類別、嚴重程度、標題、描述、改善建議），並計算 0–100 分評分。")
    add_bullet(doc, "完成 Docker Compose 容器化部署配置（nginx、web、worker、redis、db 五容器），"
        "實現開發環境與正式部署一致性。")
    add_bullet(doc, "完成 Word 格式健檢報告自動生成功能（python-docx），"
        "報告內含封面、目錄、各維度評分摘要、發現事項清單及頁面截圖。")
    add_bullet(doc, "透過初評驗收，展示系統完整掃描流程：從 URL 提交、點數扣除、"
        "BFS 爬蟲、四維掃描，到結果展示與報告下載之端對端功能演示。")
    add_body(doc,
        "最終目標為交付一套可完整運作的 Argus 網站健檢 SaaS 平台，"
        "系統支援同時管理多個掃描任務，並具備完整的點數計費稽核軌跡。", indent=True)

# ── 第2章　營運計畫 ──────────────────────────────────

def _bmc_table(doc):
    add_table_caption(doc, "表 2-2-1　商業模式畫布（Business Model Canvas）")
    tbl = doc.add_table(rows=3, cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.cell(0, 0).merge(tbl.cell(1, 0))
    tbl.cell(0, 2).merge(tbl.cell(1, 2))
    tbl.cell(0, 4).merge(tbl.cell(1, 4))
    tbl.cell(2, 0).merge(tbl.cell(2, 1))
    tbl.cell(2, 2).merge(tbl.cell(2, 4))
    bmc_data = {
        (0,0): ("① 關鍵夥伴\n(Key Partners)",
                ["雲端服務商（AWS / GCP）","支付閘道（綠界 ECPay）",
                 "AI 模型供應商（MiniMax / Google Gemini）","企業客戶（白牌合作）"], "E4DFEC"),
        (0,1): ("② 關鍵活動\n(Key Activities)",
                ["系統開發與日常維運","四維掃描引擎持續優化",
                 "AI 代理（Hermes）訓練迭代","客戶支援與教育訓練"], "DEEAF6"),
        (0,2): ("④ 價值主張\n(Value Propositions)",
                ["一站式四維網站健檢（SEO+AEO+GEO+資安）",
                 "AI 代理自動 UX 測試（Phase 2）",
                 "大幅降低工具訂閱成本（按需點數制）","Word 格式專業報告輸出"], "E2EFDA"),
        (0,3): ("⑤ 顧客關係\n(Customer Relationships)",
                ["自助式 SaaS 自動化流程","線上即時客服支援",
                 "電子郵件行銷活動","使用者社群與論壇"], "FFF2CC"),
        (0,4): ("⑦ 目標客群\n(Customer Segments)",
                ["中小企業主（電商/品牌網站）","數位行銷人員與代理商",
                 "網站開發者與自由接案者","企業 IT / 資安部門"], "FCE4EC"),
        (1,1): ("③ 關鍵資源\n(Key Resources)",
                ["掃描引擎核心技術","Hermes-Agent AI 模組",
                 "雲端基礎設施（Docker）","開發團隊人力"], "DEEAF6"),
        (1,3): ("⑥ 通路\n(Channels)",
                ["官方網站直接銷售","數位行銷代理商合作",
                 "GitHub 開源社群曝光","搜尋引擎廣告（SEA）"], "FFF2CC"),
        (2,0): ("⑧ 成本結構\n(Cost Structure)",
                ["雲端伺服器費用（AWS EC2 / RDS）","AI API 呼叫費用（MiniMax / Gemini）",
                 "開發與維護人力成本","支付手續費（ECPay）"], "F2F2F2"),
        (2,2): ("⑨ 收益流\n(Revenue Streams)",
                ["點數購買（主要，按需付費）","企業月訂閱方案（固定費用）",
                 "客製化健檢報告服務","白牌 API 授權（B2B）"], "F2F2F2"),
    }
    for (r, c), (title, items, fill) in bmc_data.items():
        cell = tbl.cell(r, c)
        _cell_shading(cell, fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p0 = cell.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p0.add_run(title)
        _set_run_font(run, Pt(10), bold=True)
        _set_para_spacing(p0, before=2, after=2, ls=1.2)
        for item in items:
            _cell_add_para(cell, item, size=Pt(9), left_indent_cm=0.3, bullet="• ")
    _keep_table_together(tbl)
    doc.add_paragraph()


def _swot_table(doc):
    add_table_caption(doc, "表 2-4-1　競爭力分析 SWOT-TOWS 矩陣")
    tbl = doc.add_table(rows=3, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    def write_hdr(cell, text, fill):
        _cell_shading(cell, fill)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _set_run_font(run, Pt(11), bold=True, color=RGBColor(0, 0, 0))
        _set_para_spacing(p, before=4, after=4, ls=1.2)
    def write_cell(cell, title, items, fill):
        _cell_shading(cell, fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p0 = cell.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p0.add_run(title); _set_run_font(run, Pt(10), bold=True)
        _set_para_spacing(p0, before=2, after=2, ls=1.2)
        for item in items:
            _cell_add_para(cell, item, size=Pt(9), left_indent_cm=0.2, bullet="➤ ")
    write_hdr(tbl.cell(0,0), "SWOT-TOWS\n分析矩陣", HDR_FILL)
    write_cell(tbl.cell(0,1), "優勢（Strengths）", [
        "S1：四維掃描差異化定位",
        "S2：AI 代理 UX 測試市場首創",
        "S3：點數制低門檻按需付費",
        "S4：Word 格式專業報告輸出",
        "S5：圖形化 GUI 操作友善",
    ], "DAEEF3")
    write_cell(tbl.cell(0,2), "劣勢（Weaknesses）", [
        "W1：品牌知名度仍處早期建立",
        "W2：AI 代理（Phase 2）待商業驗證",
        "W3：資安掃描目前僅限被動檢測",
        "W4：初期技術支援人力有限",
        "W5：AEO/GEO 市場教育成本高",
    ], "FDE9D9")
    write_cell(tbl.cell(1,0), "機會（Opportunities）", [
        "O1：生成式 AI 帶動 AEO/GEO 需求爆發",
        "O2：政府推動中小企業數位轉型補助",
        "O3：競品普遍缺乏 AEO/GEO 功能",
        "O4：企業資安合規意識提升（個資法）",
    ], "EBF1DE")
    write_cell(tbl.cell(1,1), "SO 策略（攻勢）", [
        "積極進攻 AEO/GEO 功能藍海市場（S1+O1）",
        "申請數位轉型補助計畫取得資金（S3+O2）",
        "推出 AI 代理免費試用建立品牌聲量（S2+O3）",
        "強化資安標頭掃描因應合規需求（S1+O4）",
    ], "E8F5E9")
    write_cell(tbl.cell(1,2), "WO 策略（扭轉）", [
        "與數位行銷代理商合作白牌推廣（W1+O2）",
        "參與政府補助計畫提升品牌公信力（W1+O2）",
        "開設 AEO/GEO 教育內容降低認知門檻（W5+O1）",
        "快速迭代主動式資安掃描功能（W3+O4）",
    ], "FFF9C4")
    write_cell(tbl.cell(2,0), "威脅（Threats）", [
        "T1：Google 演算法頻繁更新使規則失效",
        "T2：Ahrefs/SEMrush 等大廠可能跟進",
        "T3：AI API 成本波動影響毛利率",
        "T4：市場教育成本高（AEO/GEO 認知度低）",
    ], "F2DCDB")
    write_cell(tbl.cell(2,1), "ST 策略（多角化）", [
        "強化報告差異化深度對抗大廠（S1+T2）",
        "建立 LTV 高的企業訂閱方案（S3+T3）",
        "持續跟進 Google 演算法更新迭代（S1+T1）",
        "以 AI 代理功能建立護城河（S2+T2）",
    ], "E3F2FD")
    write_cell(tbl.cell(2,2), "WT 策略（防禦）", [
        "快速迭代縮短 Phase 2 商業驗證週期（W2+T2）",
        "建立使用者回饋機制持續改善（W4+T1）",
        "控制 API 呼叫成本優化計費模型（W4+T3）",
        "與資安社群合作強化掃描深度（W3+T2）",
    ], "FCE4EC")
    _keep_table_together(tbl)
    doc.add_paragraph()


def ch2(doc):
    add_chapter(doc, "第2章　營運計畫")
    add_section(doc, "2-1　可行性分析")
    add_body(doc, "（一）技術可行性", bold=True)
    add_body(doc,
        "本系統所採用的核心技術均為成熟的開源框架與商業 API，技術可行性高。"
        "Django 5.1 + DRF 3.15 為全球最廣泛使用的 Python Web 框架之一，"
        "擁有完善文件與社群支援（Django Software Foundation, 2024）。"
        "Playwright 1.45 由 Microsoft 官方維護，穩定性高。"
        "PostgreSQL 16 為業界標準關聯式資料庫，支援 UUID 主鍵與 JSONB 欄位。"
        "Redis 7.2 + Celery 5.4 為非同步任務佇列的成熟組合，被大型 SaaS 平台廣泛採用。",
        indent=True)
    add_std_table(doc,
        ["評估維度","技術選型","成熟度","可行性評估"],
        [
            ("後端框架","Django 5.1 + DRF 3.15","★★★★★","高：主流框架，文件完整"),
            ("前端框架","React 18 + Vite 5","★★★★★","高：社群活躍，生態完整"),
            ("爬蟲引擎","Playwright 1.45 + Chromium","★★★★☆","高：Microsoft 維護"),
            ("任務佇列","Celery 5.4 + Redis 7.2","★★★★★","高：業界標準組合"),
            ("資料庫","PostgreSQL 16","★★★★★","高：企業級穩定性"),
            ("容器化","Docker 26 + Compose 2.28","★★★★★","高：部署一致性佳"),
            ("AI API","MiniMax / Google Gemini","★★★★☆","中高：API 成本需控管"),
        ],
        "表 2-1-1　技術可行性評估表")
    add_body(doc, "（二）經濟可行性", bold=True)
    add_body(doc,
        "開發採用開源技術棧（Django、React、PostgreSQL、Celery 等），無框架授權費用，"
        "顯著降低初期投入成本。營運面採雲端部署，依實際使用量計費；"
        "計費模式為按需點數制，使用者用多少付多少，無強制月費負擔。"
        "相較於下圖所列國際 SEO 工具動輒每月數千元的訂閱費，本系統具明顯成本優勢。",
        indent=True)
    _insert_chart(doc, _make_cost_chart(), "圖 2-1-1　主要競品月費比較")
    add_body(doc,
        "資料來源：各工具官方定價頁（2026 年 5 月查得）——"
        "Google Search Console https://search.google.com/search-console/about（免費）；"
        "Screaming Frog https://www.screamingfrog.co.uk/seo-spider/pricing/（€245/年）；"
        "Ahrefs https://ahrefs.com/pricing（Starter US$29/月）；"
        "SEMrush https://www.semrush.com/pricing/seo/（Pro US$139.95/月）。"
        "外幣依台灣銀行牌告匯率（https://rate.bot.com.tw/xrt）約 1 USD≈NT$32、1 EUR≈NT$35 換算，"
        "匯率每日浮動，數值僅供比較參考。",
        indent=True)
    add_body(doc, "（三）操作可行性", bold=True)
    add_body(doc,
        "系統採用全圖形化 Web 介面，操作流程簡化為三步：輸入目標網址→選擇掃描模式→"
        "下載報告，平均學習時間短。台灣中小企業基數龐大——經濟部《2024 年中小企業白皮書》"
        "指出全台中小企業逾 167.4 萬家，占全體企業 98% 以上，其中具網站經營需求者"
        "為本系統的潛在使用族群，市場基礎充足。", indent=True)

    add_section(doc, "2-2　商業模式（Business Model）")
    add_body(doc,
        "本系統依據 Osterwalder & Pigneur（2010）提出之商業模式畫布（BMC）框架進行分析，"
        "九大要素以標準九宮格形式呈現。", indent=True)
    _bmc_table(doc)

    add_section(doc, "2-3　市場分析（STP）")
    add_body(doc,
        "本系統採用 STP 框架（Segmentation、Targeting、Positioning）"
        "系統性定義目標市場與競爭定位。", indent=True)
    add_body(doc, "➤ 市場區隔（Segmentation）", bold=True)
    add_bullet(doc, "地理區隔：以台灣及東南亞繁體中文市場為主，次要目標為日韓英語市場。台灣中小企業逾 167.4 萬家（經濟部《2024 年中小企業白皮書》），具網站經營與健檢需求者眾。")
    add_bullet(doc, "人口區隔：25–45 歲，具數位行銷或網站開發背景之專業人士；企業規模以 10–500 人之中小型企業為主。")
    add_bullet(doc, "行為區隔：曾使用過 Google Search Console 或 SEO 工具，具數據導向決策習慣，對網站效能與安全性有所關注。")
    add_bullet(doc, "心理區隔：重視效率、成本意識強，傾向訂閱制 SaaS，不願為多套工具付費，期望單一平台解決多元問題。")
    add_body(doc, "➤ 目標市場（Targeting）", bold=True)
    add_bullet(doc, "主要目標：台灣中小企業電商、品牌網站經營者。此族群對 SEO 有強烈需求，但鮮少有能力同時處理 AEO/GEO 與資安合規問題。")
    add_bullet(doc, "次要目標：數位行銷代理商（可白牌使用 Argus API 服務其客戶）、自由接案網頁開發者（需快速為客戶提供健檢報告）。")
    add_body(doc, "➤ 市場定位（Positioning）", bold=True)
    add_bullet(doc, "定位聲明：「Argus 是唯一同時涵蓋 SEO、AEO、GEO 與資安的一站式網站健檢平台，以按需點數計費模式，讓中小企業以 1/10 成本享受企業級診斷服務。」")
    add_bullet(doc, "差異化優勢一：四維評分框架——市場上唯一整合 AEO 與 GEO 的商業化工具。")
    add_bullet(doc, "差異化優勢二：按需計費——採點數制，用多少付多少，無強制月費。")
    add_bullet(doc, "差異化優勢三：圖形化操作——相較於 Nmap、Dirsearch、Katana 等命令列工具，Argus 提供完整 GUI，大幅降低操作門檻。")

    add_section(doc, "2-4　競爭力分析（SWOT-TOWS 分析）")
    add_body(doc,
        "本節採用 SWOT-TOWS 整合分析矩陣，同時呈現 Argus 的優劣勢、市場機會與威脅，"
        "以及對應的四種策略方向（SO 攻勢、WO 扭轉、ST 多角化、WT 防禦）。"
        "矩陣垂直維度區分企業內部（優勢/劣勢）與企業外部（機會/威脅），"
        "水平維度區分有利與不利因素。", indent=True)
    _swot_table(doc)


# ── 第3章　系統規格 ──────────────────────────────────

def ch3(doc):
    add_chapter(doc, "第3章　系統規格")

    add_section(doc, "3-1　系統架構")
    add_body(doc,
        "Argus 系統採用現代化 SaaS 分層架構，由下而上依序為：用戶端層、反向代理層、"
        "應用程式層、非同步任務層及資料持久層。各層之間透過 HTTP/REST、WebSocket 及"
        "Redis 訊息佇列通訊，確保各元件職責清晰且可獨立水平擴展。"
        "系統整體架構如圖 3-1-1 所示（PlantUML 圖稿請見附件 plantuml_diagrams_v3.txt）。",
        indent=True)
    add_placeholder(doc, "圖 3-1-1　Argus 系統架構圖", width=Cm(14))
    add_body(doc, "各層的主要職責說明如下：")
    add_bullet(doc, "用戶端層（Client Layer）：使用者透過 Chrome/Firefox/Edge 瀏覽器存取 React 18 SPA，"
        "無需安裝任何客戶端軟體。")
    add_bullet(doc, "反向代理層（Nginx 1.26）：處理 SSL 終止（HTTPS 443→HTTP 80）、"
        "靜態資源服務（React dist/）及後端 API 反向代理（/api/* → Django 8000）。")
    add_bullet(doc, "應用程式層（Django 5.1）：提供 RESTful API 端點（Django REST Framework 3.15）、"
        "使用者認證（Google OAuth 2.0）、點數計費（BillingService）及 Word 報告生成（python-docx）。")
    add_bullet(doc, "非同步任務層（Celery 5.4 + Redis 7.2）：接收掃描任務請求，執行 BFS 爬蟲"
        "（Playwright Chromium）及四維掃描引擎（SEO/AEO/GEO/Security），"
        "掃描結果寫回 PostgreSQL。")
    add_bullet(doc, "資料持久層（PostgreSQL 16）：儲存所有業務資料，"
        "包含使用者、掃描任務、頁面、發現事項、點數交易及評論等 10 個資料表。")

    add_section(doc, "3-2　系統軟、硬體需求與技術平台")
    add_body(doc, "（一）硬體環境需求", bold=True)
    add_body(doc,
        "本節僅列出系統正式上線之公網部署環境與終端使用者之用戶端需求；"
        "開發測試環境非屬正式營運配置，故不納入。", indent=True)
    add_std_table(doc,
        ["環境","類別","規格需求","備註"],
        [
            ("正式部署","CPU","vCPU × 4（AWS t3.xlarge 或同等級）","Celery Worker 並行需求"),
            ("正式部署","RAM","16 GB","Django + Celery + PostgreSQL + Redis"),
            ("正式部署","儲存空間","SSD 100 GB（含資料庫與截圖儲存）","截圖約 50–200 KB/頁"),
            ("正式部署","頻寬","100 Mbps（對外爬蟲流量需求）","BFS 爬蟲並發需求"),
            ("用戶端（瀏覽器）","瀏覽器","Chrome 120+、Firefox 120+、Edge 120+","支援 ES2020"),
        ],
        "表 3-2-1　硬體環境需求規格")

    add_body(doc, "（二）軟體環境需求與技術平台", bold=True)
    add_std_table(doc,
        ["類別","技術/版本","用途說明","官方文件"],
        [
            ("後端語言","Python 3.12.x","Django ORM / Celery 任務佇列","https://docs.python.org/3.12/"),
            ("Web 框架","Django 5.1","RESTful API 與 Admin 後台","https://docs.djangoproject.com/en/5.1/"),
            ("API 框架","Django REST Framework 3.15","Serializer、ViewSet、Permission","https://www.django-rest-framework.org/"),
            ("前端框架","React 18 + Vite 5","SPA 單頁應用程式","https://react.dev/ / https://vitejs.dev/"),
            ("UI 元件庫","Tailwind CSS 3.4 + shadcn/ui","響應式設計元件","https://tailwindcss.com/"),
            ("視覺化","ReactFlow 11","網站拓樸圖（節點流程圖）","https://reactflow.dev/"),
            ("資料庫","PostgreSQL 16","主要關聯式資料庫","https://www.postgresql.org/docs/16/"),
            ("訊息佇列","Redis 7.2","Celery Broker + 結果快取","https://redis.io/docs/"),
            ("非同步任務","Celery 5.4","背景掃描任務調度","https://docs.celeryq.dev/en/stable/"),
            ("爬蟲引擎","Playwright 1.45 + Chromium","無頭瀏覽器 BFS 爬蟲","https://playwright.dev/"),
            ("容器化","Docker 26 + Compose 2.28","開發與正式部署一致性","https://docs.docker.com/"),
            ("反向代理","Nginx 1.26","靜態資源服務 + SSL 終止","https://nginx.org/en/docs/"),
            ("套件管理","uv 0.4（Python）/ npm 10（Node）","環境隔離與版本鎖定","https://docs.astral.sh/uv/"),
            ("版本控制","Git 2.45 + GitHub","協作開發與版本管理","https://git-scm.com/doc"),
            ("AI API","MiniMax-M2.7 / Google Gemini","Hermes-Agent 工具調用（Phase 2）","https://platform.minimaxi.com/"),
            ("報告生成","python-docx 1.1","Word 格式健檢報告自動生成","https://python-docx.readthedocs.io/"),
            ("認證","Google OAuth 2.0","社群登入驗證","https://developers.google.com/identity"),
        ],
        "表 3-2-2　軟體環境需求規格")

    add_section(doc, "3-3　使用標準與工具")
    add_body(doc,
        "本專題在系統分析與設計過程中採用 UML 2.x（Unified Modeling Language）作為"
        "主要的軟體工程建模規範，涵蓋使用個案圖、活動圖、循序圖、類別圖、"
        "佈署圖、套件圖、元件圖及狀態機圖共八類圖形。"
        "所有 UML 圖使用 PlantUML（https://plantuml.com/）語法撰寫，"
        "圖碼見附件 plantuml_diagrams_v3.txt。使用標準與工具彙整如表 3-3-1。",
        indent=True)
    add_std_table(doc,
        ["工具 / 標準","版本 / 規格","用途","官方連結"],
        [
            ("UML 2.x","UML 2.5.1","系統分析與設計建模規範（九類圖）","https://www.omg.org/spec/UML/"),
            ("PlantUML","1.2024.x","UML 圖碼生成工具","https://plantuml.com/"),
            ("Python uv","0.4.x","Python 虛擬環境管理與套件安裝","https://docs.astral.sh/uv/"),
            ("Git / GitHub","2.45 / Actions","版本控制與多人協作開發","https://git-scm.com/"),
            ("Docker / Compose","26 / 2.28","容器化部署，確保環境一致性","https://docs.docker.com/"),
            ("Visual Studio Code","1.9x","主要整合開發環境（IDE）","https://code.visualstudio.com/"),
            ("python-docx","1.1","Word 格式健檢報告自動生成","https://python-docx.readthedocs.io/"),
            ("Matplotlib","3.9","可行性分析圖表生成","https://matplotlib.org/"),
            ("Postman","11.x","RESTful API 端點測試","https://www.postman.com/"),
            ("pgAdmin","8.x","PostgreSQL 資料庫管理","https://www.pgadmin.org/"),
        ],
        "表 3-3-1　使用標準與工具表")


# ── 第4章　專案時程與組織分工 ──────────────────────────

def _gantt_table(doc):
    """甘特圖：每個任務兩列（預期/實際），任務名稱跨列合併"""
    add_table_caption(doc, "表 4-1-1　專案甘特圖（114/09 – 115/06）")
    months = ["09月","10月","11月","12月","01月","02月","03月","04月","05月","06月"]
    tasks = [
        # (名稱, 預期月份索引list, 實際月份索引list)
        ("需求訪談與市場分析",     [0,1],         [0,1]),
        ("系統架構設計",           [1,2],         [1,2]),
        ("資料庫設計",             [2,3],         [2,3]),
        ("後端 API 開發",          [2,3,4,5],     [2,3,4,5]),
        ("BFS 爬蟲引擎實作",       [3,4,5],       [3,4,5]),
        ("四維掃描引擎開發",       [4,5,6],       [4,5,6]),
        ("前端 React UI 開發",     [3,4,5,6],     [3,4,5,6]),
        ("Celery 任務佇列整合",    [4,5],         [4,5]),
        ("Docker 容器化部署",      [5,6],         [5,6]),
        ("系統整合測試",           [6,7],         [6,7]),
        ("Hermes-Agent（Phase 2）",[7,8],         []),
        ("文件撰寫與簡報製作",     [5,6,7,8,9],   [5,6,7,8]),
        ("期末展示與驗收",         [9],           []),
    ]
    n_tasks = len(tasks)
    HDR_ROWS = 2                     # 年份列 + 月份列
    n_rows  = HDR_ROWS + n_tasks * 2 + 1   # 雙表頭 + task*2 + legend
    n_cols  = 1 + len(months)        # 任務名稱 + 10 個月
    # 114 年（09–12 月，4 欄）／115 年（01–06 月，6 欄）
    year_spans = [("中華民國 114 年", 0, 3), ("中華民國 115 年", 4, 9)]

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ── 年份列（row 0）：左上角與月份列垂直合併 ──
    tbl.cell(0, 0).merge(tbl.cell(1, 0))
    c_corner = tbl.cell(0, 0)
    _cell_write(c_corner, "工作項目", size=Pt(10), bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_shading(c_corner, HDR_FILL)
    c_corner.paragraphs[0].runs[0].font.color.rgb = RGBColor(0,0,0)
    for label, c0, c1 in year_spans:
        tbl.cell(0, c0 + 1).merge(tbl.cell(0, c1 + 1))
        c = tbl.cell(0, c0 + 1)
        _cell_write(c, label, size=Pt(9), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_shading(c, HDR_FILL)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0,0,0)

    # ── 月份列（row 1）──
    hrow = tbl.rows[1]
    for mi, m in enumerate(months):
        c = hrow.cells[mi + 1]
        _cell_write(c, m, size=Pt(9), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_shading(c, HDR_FILL)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0,0,0)

    # ── 任務列 ──
    for ti, (name, exp, act) in enumerate(tasks):
        row_a = HDR_ROWS + ti * 2
        row_b = row_a + 1
        # 合併任務名稱跨兩列
        tbl.cell(row_a, 0).merge(tbl.cell(row_b, 0))
        c_name = tbl.cell(row_a, 0)
        _cell_write(c_name, name, size=Pt(9), align=WD_ALIGN_PARAGRAPH.LEFT)
        _cell_shading(c_name, "FFFFFF")
        # 預期列（淺灰）
        for mi in range(len(months)):
            c = tbl.cell(row_a, mi + 1)
            _cell_shading(c, GANTT_EXPECTED if mi in exp else "FFFFFF")
        # 實際列（深灰）
        for mi in range(len(months)):
            c = tbl.cell(row_b, mi + 1)
            _cell_shading(c, GANTT_ACTUAL if mi in act else "FFFFFF")

    # ── 圖例列 ──
    leg_row = tbl.rows[n_rows - 1]
    tbl.cell(n_rows-1, 0).merge(tbl.cell(n_rows-1, n_cols-1))
    c_leg = tbl.cell(n_rows-1, 0)
    _cell_shading(c_leg, "F5F5F5")
    p = c_leg.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run("  ■ ")
    r1.font.color.rgb = RGBColor(0xD9,0xD9,0xD9); r1.font.size = Pt(14)
    r2 = p.add_run(" 預期進度　")
    _set_run_font(r2, Pt(10))
    r3 = p.add_run("■ ")
    r3.font.color.rgb = RGBColor(0x59,0x59,0x59); r3.font.size = Pt(14)
    r4 = p.add_run(" 實際進度")
    _set_run_font(r4, Pt(10))

    _table_fixed_widths(tbl, [3.2] + [1.1] * len(months))
    _table_no_split(tbl, repeat_header=True, header_rows=HDR_ROWS)
    doc.add_paragraph()


# 4 名成員 × 專業分組（後端/前端/美術/文件），● 主要、○ 次要
DIVISION_MEMBERS = [
    "○○○\n（後端開發）", "○○○\n（前端開發）",
    "○○○\n（美術設計）", "○○○\n（文件撰寫統整）",
]
DIVISION_GROUPS = [
    ("後端開發", [
        ("後端 API 開發（Django + DRF）",      ["●", "", "", "○"]),
        ("PostgreSQL 資料庫建置",              ["●", "", "", "○"]),
        ("點數計費系統（BillingService）",     ["●", "○", "", ""]),
        ("BFS 爬蟲與四維掃描引擎",             ["●", "", "○", ""]),
    ]),
    ("前端開發", [
        ("React 18 SPA 介面開發",              ["", "●", "○", ""]),
        ("ReactFlow 網站拓樸圖",               ["○", "●", "", ""]),
        ("後台管理介面開發",                   ["", "●", "", "○"]),
    ]),
    ("美術設計", [
        ("UI/UX 視覺設計",                     ["", "○", "●", ""]),
        ("品牌識別與配色風格",                 ["", "", "●", "○"]),
        ("簡報與展示視覺設計",                 ["", "", "●", "○"]),
    ]),
    ("文件撰寫統整", [
        ("系統手冊撰寫統整",                   ["○", "", "", "●"]),
        ("GitHub 紀錄與版本管理",              ["○", "", "", "●"]),
        ("期末簡報製作",                       ["", "", "○", "●"]),
    ]),
]


def _division_table(doc):
    """專業組織分工表：4 名成員，列依後端/前端/美術/文件分組；● 主要、○ 次要"""
    add_table_caption(doc, "表 4-2-1　專案組織分工表（● 主要負責，○ 次要協助）")
    n_items = sum(len(items) for _, items in DIVISION_GROUPS)
    n_rows = 1 + len(DIVISION_GROUPS) + n_items  # 標題 + 分組列 + 工作列
    n_cols = 1 + len(DIVISION_MEMBERS)
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 標題列
    _cell_write(tbl.cell(0, 0), "專業分組 / 工作項目", size=Pt(10), bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_shading(tbl.cell(0, 0), HDR_FILL)
    tbl.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
    for mi, m in enumerate(DIVISION_MEMBERS):
        c = tbl.cell(0, mi + 1)
        _cell_write(c, m, size=Pt(9), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_shading(c, HDR_FILL)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)

    r = 1
    for gname, items in DIVISION_GROUPS:
        # 分組標題列（整列合併、淺色）
        tbl.cell(r, 0).merge(tbl.cell(r, n_cols - 1))
        gc = tbl.cell(r, 0)
        _cell_write(gc, f"【{gname}】", size=Pt(10), bold=True,
                    align=WD_ALIGN_PARAGRAPH.LEFT)
        _cell_shading(gc, "EDEDED")
        r += 1
        for ti, (item, marks) in enumerate(items):
            _cell_write(tbl.cell(r, 0), item, size=Pt(9), align=WD_ALIGN_PARAGRAPH.LEFT)
            _cell_shading(tbl.cell(r, 0), ALT_FILL if ti % 2 == 0 else "FFFFFF")
            for mi, mk in enumerate(marks):
                c = tbl.cell(r, mi + 1)
                _cell_write(c, mk, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
                _cell_shading(c, ALT_FILL if ti % 2 == 0 else "FFFFFF")
            r += 1

    _table_fixed_widths(tbl, [5.0] + [2.3] * len(DIVISION_MEMBERS))
    _table_no_split(tbl, repeat_header=True)
    doc.add_paragraph()


def ch4(doc):
    add_chapter(doc, "第4章　專案時程與組織分工")

    add_section(doc, "4-1　專案時程")
    add_body(doc,
        "本專題自 114 年 9 月起始，預計於 115 年 6 月完成複評，"
        "整體開發週期約 10 個月。甘特圖以淺灰色（■）表示預期進度，"
        "深灰色（■）表示實際進度，每一工作項目分為上下兩列呈現。", indent=True)
    _gantt_table(doc)

    add_section(doc, "4-2　專案組織與分工")
    add_body(doc,
        "本組共 4 名成員，依專業分為後端開發、前端開發、美術設計、文件撰寫統整四組。"
        "每項工作指定 1 位主要負責人（●），次要協助最多 2 位（○），"
        "每一項只能有 1 位主要負責人。分工表如表 4-2-1 所示。", indent=True)
    _division_table(doc)

    add_body(doc, "各組員工作內容與貢獻度如表 4-2-2 所示：")
    add_std_table(doc,
        ["序號","姓名","工作內容（各限 100 字以內）","貢獻度"],
        [
            ("1","○○○（後端開發）",
             "負責整體系統架構與後端 API 開發（Django + DRF）、PostgreSQL 資料庫建置、"
             "點數計費系統（BillingService 原子交易機制）及 BFS 爬蟲與四維掃描引擎核心邏輯。","25%"),
            ("2","○○○（前端開發）",
             "負責 React 18 前端 SPA 開發（掃描提交、結果展示）、ReactFlow 網站拓樸圖、"
             "後台管理介面開發及前端與後端 API 串接整合。","25%"),
            ("3","○○○（美術設計）",
             "負責全站 UI/UX 視覺設計、品牌識別與配色風格制定、"
             "簡報與展示視覺設計，並協助前端介面切版與美術資源產出。","25%"),
            ("4","○○○（文件撰寫統整）",
             "負責系統手冊撰寫與統整、GitHub 紀錄與版本管理、期末簡報製作，"
             "並協助系統整合測試與文件校對。","25%"),
        ],
        "表 4-2-2　專題成果工作內容與貢獻度表")

    add_section(doc, "4-3　上傳 GitHub 紀錄")
    add_body(doc,
        "本專題原始碼統一於 GitHub 進行版本控制，採用 Git Flow 工作流程管理，"
        "主要分支說明如下：", indent=True)
    add_bullet(doc, "main：正式部署分支，僅接受 Pull Request 合併，受保護分支")
    add_bullet(doc, "develop：主要開發整合分支，功能開發完成後 PR 至 develop")
    add_bullet(doc, "feature/*：各功能開發分支，完成後 PR 至 develop")
    add_bullet(doc, "hotfix/*：緊急修復分支，直接合併至 main 及 develop")
    add_body(doc, "GitHub Repository：https://github.com/○○○/argus（請替換為實際網址）")
    add_body(doc, "以下圖片為各組員的 GitHub 提交紀錄截圖，請於初評前各截一次：")
    add_placeholder(doc, "圖 4-3-1　GitHub 提交紀錄截圖", width=Cm(14))


# ── 第5章　需求模型 ──────────────────────────────────

def ch5(doc):
    add_chapter(doc, "第5章　需求模型")
    add_body(doc, "（系統分析與設計使用 UML 2.x 建模，所有 UML 圖碼見 plantuml_diagrams_v3.txt）")

    add_section(doc, "5-1　使用者需求")
    add_body(doc, "（一）功能需求（Functional Requirements）", bold=True)
    add_std_table(doc,
        ["編號","需求名稱","說明","優先度"],
        [
            ("FR-01","使用者登入","支援 Google OAuth 2.0 社群登入及開發模式帳密登入","高"),
            ("FR-02","提交掃描任務","輸入目標 URL，選擇掃描模式（passive/active）與最大頁數","高"),
            ("FR-03","SEO 掃描","檢測 Meta 標籤、標題結構、圖片 alt、robots.txt、sitemap、頁面速度","高"),
            ("FR-04","AEO 掃描","檢測結構化資料（Schema.org JSON-LD）、FAQ/HowTo 標記、精選摘要適配性","高"),
            ("FR-05","GEO 掃描","評估語意密度、長尾關鍵字覆蓋率、AI 引擎可讀性評分","高"),
            ("FR-06","資安被動掃描","檢測 HTTPS/HSTS/CSP/XFO/XSS-Protection 等 HTTP 安全標頭","高"),
            ("FR-07","BFS 網頁爬蟲","以廣度優先搜尋爬取所有可達頁面，支援 max_depth / max_pages 設定","高"),
            ("FR-08","點數計費","hold→settle→refund 原子交易機制，支援月贈點與方案購買","高"),
            ("FR-09","下載健檢報告","掃描完成後自動生成 .docx 格式專業報告供下載","高"),
            ("FR-10","ReactFlow 拓樸圖","視覺化呈現 BFS 爬蟲涵蓋的頁面節點結構","中"),
            ("FR-11","管理員後台","Staff 可管理使用者點數、查看掃描任務、回覆評論","中"),
            ("FR-12","Hermes-Agent（Phase 2）","AI 代理以 observe→think→act 循環自動 UX 測試","低"),
        ],
        "表 5-1-1　功能需求清單")

    add_body(doc, "（二）非功能需求（Non-Functional Requirements）", bold=True)
    add_std_table(doc,
        ["編號","需求類別","需求說明","衡量指標"],
        [
            ("NFR-01","效能","掃描結果 API 回應時間 ≤ 500ms（P95）","Locust 壓測 100 同時用戶"),
            ("NFR-02","可靠性","系統可用性 ≥ 99%（月計算）","Docker health check 自動重啟"),
            ("NFR-03","安全性","所有 API 端點須通過 JWT 或 Session 認證","Django IsAuthenticated 防護"),
            ("NFR-04","可維護性","後端測試覆蓋率 ≥ 80%（Django TestCase）","192 項自動化測試"),
            ("NFR-05","可擴展性","Celery Worker 支援水平擴展（多 Worker 並行）","Docker Compose scale"),
            ("NFR-06","資料一致性","計費操作使用 select_for_update + transaction.atomic","BillingService 冪等保護"),
            ("NFR-07","隱私合規","API 端點不得直接回傳使用者個資，需透過 Serializer whitelist","DRF Serializer 實施"),
        ],
        "表 5-1-2　非功能需求清單")

    add_section(doc, "5-2　使用個案圖（Use Case Diagram）")
    add_body(doc,
        "系統定義三類參與者（Actor）：一般使用者（User）、管理員 Staff（Staff）"
        "及 Celery Worker（系統背景服務）。主要使用個案如圖 5-2-1 所示，"
        "涵蓋登入、提交掃描、查看結果、下載報告、購買點數、撰寫評論等 12 個案例。",
        indent=True)
    add_placeholder(doc, "圖 5-2-1　使用個案圖（Use Case Diagram）", width=Cm(14))

    add_section(doc, "5-3　使用個案描述（Activity Diagram）")
    add_body(doc,
        "以下以提交掃描任務為主要使用個案，以活動圖（Activity Diagram）描述"
        "從使用者輸入網址到系統完成掃描的完整流程，涵蓋輸入驗證、點數檢查、"
        "Celery 任務派發、BFS 爬蟲、四維掃描及報告生成六個主要階段，"
        "並標示失敗與取消路徑。", indent=True)
    add_placeholder(doc, "圖 5-3-1　活動圖：提交掃描任務（Activity Diagram）", width=Cm(14))
    add_std_table(doc,
        ["欄位","說明"],
        [
            ("使用個案名稱","提交掃描任務（Submit Scan Job）"),
            ("主要參與者","一般使用者（User）"),
            ("前置條件","使用者已登入，且錢包點數 ≥ 所需預扣額度"),
            ("正常流程",
             "1. 使用者輸入目標網址與掃描設定（模式、最大頁數）\n"
             "2. 系統驗證 URL 格式合法性\n"
             "3. BillingService 執行 hold_for_scan（預扣點數）\n"
             "4. 建立 ScanJob 記錄（status=queued）\n"
             "5. Celery 派發背景任務\n"
             "6. Worker 執行 BFS 爬蟲（status=crawling）\n"
             "7. Worker 逐頁執行四維掃描（status=scanning）\n"
             "8. 掃描完成，settle_scan_actual 結算實際點數\n"
             "9. 自動生成 Word 報告（status=completed）"),
            ("替代流程",
             "A1. URL 格式錯誤 → 回傳 HTTP 400\n"
             "A2. 點數不足 → 回傳 HTTP 402，引導至購買頁面\n"
             "A3. 使用者取消 → cancel API 呼叫，refund_full_for_scan 全退"),
            ("後置條件","ScanJob.status=completed，Finding 已寫入資料庫，Word 報告可下載"),
        ],
        "表 5-3-1　使用個案描述：提交掃描任務")

    add_section(doc, "5-4　分析類別圖（Analysis Class Diagram）")
    add_body(doc,
        "圖 5-4-1 為系統分析層級的類別圖，展示八個核心實體——"
        "User、ScanJob、Finding、Page、CoinWallet、CoinTransaction、"
        "PurchaseOrder 及 AgentSession（Phase 2，虛線框標示）——之間的關聯關係。",
        indent=True)
    add_placeholder(doc, "圖 5-4-1　分析類別圖（Analysis Class Diagram）", width=Cm(14))


# ── 第6章　設計模型 ──────────────────────────────────

def ch6(doc):
    add_chapter(doc, "第6章　設計模型")

    add_section(doc, "6-1　循序圖（Sequential Diagram）")
    add_body(doc,
        "圖 6-1-1 為提交掃描任務之循序圖，涉及六個生命線："
        "使用者瀏覽器（User Browser）、Django REST API（DRF）、"
        "計費服務（BillingService）、Celery 任務佇列（Celery Queue）、"
        "Celery Worker 及資料庫（PostgreSQL）。"
        "核心流程包含預扣點數（hold_for_scan）、任務入佇列、BFS 爬蟲、"
        "四維掃描、結算點數（settle_scan_actual）及 Word 報告生成共 18 個互動步驟。",
        indent=True)
    add_placeholder(doc, "圖 6-1-1　掃描任務循序圖（Sequential Diagram）", width=Cm(14))

    add_section(doc, "6-2　設計類別圖（Design Class Diagram）")
    add_body(doc,
        "圖 6-2-1 展示系統設計層級的類別圖，每個類別包含三個區塊：名稱、屬性及方法。"
        "BillingService 以 «service» 構造型標示，採用靜態方法設計，"
        "是 CoinWallet 與 CoinTransaction 的唯一合法寫入入口，"
        "所有計費操作透過 select_for_update + transaction.atomic 保護。",
        indent=True)
    add_placeholder(doc, "圖 6-2-1　設計類別圖（Design Class Diagram）", width=Cm(14))
    add_std_table(doc,
        ["類別名稱","主要屬性","主要方法","說明"],
        [
            ("User","id, username, email,\nis_staff, is_active","get_wallet(), is_staff_user()","繼承 Django AbstractUser"),
            ("ScanJob","id(UUID), user, original_url,\nstatus, scan_mode, max_pages,\nprogress(JSON), overall_score","cancel(), get_findings_summary()","狀態機：queued→crawling→scanning→completed"),
            ("CoinWallet","id, user, balance,\ntotal_purchased_ntd,\nlast_bonus_year/month","get_balance(), check_sufficient()","一對一對應 User，禁止直接 .save()"),
            ("CoinTransaction","id, wallet, amount, type,\nscan(FK nullable), note","（唯讀）","amount 正=入帳，負=扣款"),
            ("BillingService","（靜態類別，無實例屬性）",
             "hold_for_scan(scan, max_pages)\nsettle_scan_actual(scan, actual_pages)\nrefund_full_for_scan(scan)\ngrant_monthly_bonus(user)",
             "所有寫入操作的唯一入口，select_for_update 保護"),
            ("Finding","id, page, category,\nseverity, title,\ndescription, recommendation","（查詢用）","category: SEO/AEO/GEO/Security"),
            ("Page","id, scan, url,\nstatus_code, title, depth","（查詢用）","BFS 爬取的每一頁面"),
        ],
        "表 6-2-1　設計類別主要屬性與方法規格")


# ── 第7章　實作模型 ──────────────────────────────────

def ch7(doc):
    add_chapter(doc, "第7章　實作模型")

    add_section(doc, "7-1　佈署圖（Deployment Diagram）")
    add_body(doc,
        "Argus 採用 Docker Compose 進行容器化部署，如圖 7-1-1 所示，"
        "由五個容器組成：nginx（反向代理，port 80/443）、web（Django + React SPA，port 8000）、"
        "worker（Celery 背景任務處理器）、redis（任務佇列 Broker，port 6379）"
        "及 db（PostgreSQL，port 5432）。所有容器透過 Docker Bridge Network 互通，"
        "外部流量統一由 nginx 路由。正式環境另透過 Cloudflare Tunnel 提供 HTTPS 加密通道。",
        indent=True)
    add_placeholder(doc, "圖 7-1-1　Docker Compose 佈署圖（Deployment Diagram）", width=Cm(14))
    add_std_table(doc,
        ["容器名稱","映像","對外埠","對內通訊","職責"],
        [
            ("nginx","nginx:1.26-alpine","80, 443","web:8000","SSL 終止、靜態資源、反向代理"),
            ("web","python:3.12-slim","8000（內部）","db:5432, redis:6379","Django + DRF + React SPA"),
            ("worker","（同 web 映像）","無","db:5432, redis:6379","Celery Worker，執行掃描任務"),
            ("redis","redis:7.2-alpine","6379（內部）","worker, web","Celery Broker + 結果快取"),
            ("db","postgres:16-alpine","5432（內部）","web, worker","PostgreSQL 主資料庫"),
        ],
        "表 7-1-1　容器通訊規格")

    add_section(doc, "7-2　套件圖（Package Diagram）")
    add_body(doc,
        "圖 7-2-1 為 Django 後端的套件架構圖，展示 backend/apps/ 下七個 Django App 的依賴關係。"
        "各 App 職責清晰分離，所有 App 共用 config（Django 設定與路由）作為根套件。",
        indent=True)
    add_placeholder(doc, "圖 7-2-1　套件架構圖（Package Diagram）", width=Cm(14))

    add_section(doc, "7-3　元件圖（Component Diagram）")
    add_body(doc,
        "圖 7-3-1 以 UML 元件圖呈現系統各主要元件及其介面。"
        "元件包含 Browser（客戶端）、Nginx（:443）、Django Application（:8000）、"
        "Celery Worker、Redis Broker（:6379）、PostgreSQL（:5432）"
        "及 Playwright Chromium（虛線，外部爬蟲目標）。",
        indent=True)
    add_placeholder(doc, "圖 7-3-1　系統元件圖（Component Diagram）", width=Cm(14))

    add_section(doc, "7-4　狀態機（State Machine）")
    add_body(doc,
        "圖 7-4-1 為 ScanJob 掃描任務的狀態機圖，描述任務從建立到終止的完整生命週期。"
        "在任何階段均可轉入 failed（失敗）或 cancelled（已取消）終止狀態。",
        indent=True)
    add_placeholder(doc, "圖 7-4-1　ScanJob 狀態機圖（State Machine）", width=Cm(14))
    add_std_table(doc,
        ["狀態","觸發條件","說明"],
        [
            ("queued","ScanJob 建立後立即","任務已排入 Celery 佇列，等待 Worker 處理"),
            ("crawling","Worker 開始執行","Playwright BFS 爬蟲啟動，爬取可達頁面"),
            ("scanning","爬蟲階段完成","逐頁執行四維掃描引擎（SEO/AEO/GEO/Security）"),
            ("agent_testing","scanning 完成（Phase 2）","Hermes-Agent AI 代理執行 UX 自動測試"),
            ("completed","所有掃描階段完成","結算點數（settle_scan_actual），產生 Word 報告"),
            ("failed","任何階段發生例外","記錄錯誤訊息，執行 refund_full_for_scan 全退點數"),
            ("cancelled","使用者呼叫 cancel API","Worker 協作式取消，執行 refund_full_for_scan 全退"),
        ],
        "表 7-4-1　ScanJob 狀態說明")


# ── 第8章　資料庫設計 ────────────────────────────────

def ch8(doc):
    add_chapter(doc, "第8章　資料庫設計")

    add_section(doc, "8-1　資料庫關聯表")
    add_body(doc,
        "Argus 系統採用 PostgreSQL 16 作為主要關聯式資料庫，共包含 10 個資料表，"
        "對應 Django 的 7 個 App 模型。核心實體為 auth_user，"
        "向左連接計費系統（billing_coinwallet、billing_cointransaction、"
        "billing_pricingplan、billing_purchaseorder），"
        "向右連接掃描系統（scans_scanjob、scans_page、scans_finding），"
        "向下連接評論（reviews_platformreview）與稽核（admin_api_adminauditlog）。"
        "ER 圖如圖 8-1-1 所示（PlantUML 圖碼見 plantuml_diagrams_v3.txt）。",
        indent=True)
    add_placeholder(doc, "圖 8-1-1　資料庫 ER 圖（Entity-Relationship Diagram）", width=Cm(14))

    add_section(doc, "8-2　表格及其 Meta Data")
    add_body(doc, "以下依序說明系統 10 個資料表的欄位定義與限制（Constraints）。")

    tables_meta = [
        ("表 8-2-1　auth_user 資料表",
         ["欄位名稱","資料型別","限制","說明"],
         [
            ("id","SERIAL","PK, NOT NULL","使用者唯一識別碼（自動遞增）"),
            ("username","VARCHAR(150)","UNIQUE, NOT NULL","登入帳號"),
            ("email","VARCHAR(254)","UNIQUE, NOT NULL","電子信箱"),
            ("first_name","VARCHAR(150)","NOT NULL, DEFAULT ''","名字"),
            ("last_name","VARCHAR(150)","NOT NULL, DEFAULT ''","姓氏"),
            ("is_staff","BOOLEAN","NOT NULL, DEFAULT FALSE","Staff 管理員旗標"),
            ("is_active","BOOLEAN","NOT NULL, DEFAULT TRUE","帳號啟用狀態"),
            ("is_superuser","BOOLEAN","NOT NULL, DEFAULT FALSE","超級管理員旗標"),
            ("date_joined","TIMESTAMPTZ","NOT NULL","帳號建立時間"),
            ("last_login","TIMESTAMPTZ","NULL","最近登入時間"),
         ]),
        ("表 8-2-2　billing_coinwallet 資料表",
         ["欄位名稱","資料型別","限制","說明"],
         [
            ("id","SERIAL","PK, NOT NULL","錢包唯一識別碼"),
            ("user_id","INTEGER","FK→auth_user, UNIQUE, NOT NULL","關聯使用者（一對一）"),
            ("balance","INTEGER","NOT NULL, DEFAULT 0, CHECK(≥0)","目前點數餘額"),
            ("total_purchased_ntd","DECIMAL(10,2)","NOT NULL, DEFAULT 0","累計購買金額（NT$）"),
            ("total_scans_used","INTEGER","NOT NULL, DEFAULT 0","累計掃描次數"),
            ("last_bonus_year","INTEGER","NULL","最近月贈點年份（冪等鎖）"),
            ("last_bonus_month","INTEGER","NULL","最近月贈點月份（冪等鎖）"),
            ("created_at","TIMESTAMPTZ","NOT NULL","建立時間"),
         ]),
        ("表 8-2-3　billing_cointransaction 資料表",
         ["欄位名稱","資料型別","限制","說明"],
         [
            ("id","SERIAL","PK, NOT NULL","交易唯一識別碼"),
            ("wallet_id","INTEGER","FK→billing_coinwallet, NOT NULL","關聯錢包"),
            ("amount","INTEGER","NOT NULL","交易點數（正=入帳，負=扣款）"),
            ("type","VARCHAR(20)","NOT NULL","枚舉：purchase/hold/settle/refund/bonus/manual"),
            ("scan_id","UUID","FK→scans_scanjob, NULL","關聯掃描任務（可為 NULL）"),
            ("note","TEXT","NOT NULL, DEFAULT ''","備註說明"),
            ("created_at","TIMESTAMPTZ","NOT NULL","交易時間"),
         ]),
        ("表 8-2-4　billing_pricingplan 資料表",
         ["欄位名稱","資料型別","限制","說明"],
         [
            ("id","SERIAL","PK, NOT NULL","方案唯一識別碼"),
            ("name","VARCHAR(100)","NOT NULL","方案名稱（如：基本方案 100 點）"),
            ("coins","INTEGER","NOT NULL, CHECK(>0)","購買所得點數"),
            ("price_ntd","DECIMAL(8,0)","NOT NULL, CHECK(>0)","定價（新台幣）"),
            ("is_active","BOOLEAN","NOT NULL, DEFAULT TRUE","是否上架中"),
            ("created_at","TIMESTAMPTZ","NOT NULL","建立時間"),
         ]),
        ("表 8-2-5　billing_purchaseorder 資料表",
         ["欄位名稱","資料型別","限制","說明"],
         [
            ("id","UUID","PK, NOT NULL","訂單唯一識別碼（UUID 防爬蟲）"),
            ("user_id","INTEGER","FK→auth_user, NOT NULL","關聯使用者"),
            ("plan_id","INTEGER","FK→billing_pricingplan, NOT NULL","關聯購買方案"),
            ("transaction_id","INTEGER","FK→billing_cointransaction, NULL","關聯點數入帳交易"),
            ("status","VARCHAR(20)","NOT NULL","枚舉：pending/paid/failed/refunded"),
            ("created_at","TIMESTAMPTZ","NOT NULL","建立時間"),
            ("paid_at","TIMESTAMPTZ","NULL","付款完成時間"),
         ]),
        ("表 8-2-6　scans_scanjob 資料表",
         ["欄位名稱","資料型別","限制","說明"],
         [
            ("id","UUID","PK, NOT NULL","任務唯一識別碼（UUID 防爬蟲探測）"),
            ("user_id","INTEGER","FK→auth_user, NOT NULL","關聯使用者"),
            ("original_url","VARCHAR(2048)","NOT NULL","目標網站 URL"),
            ("status","VARCHAR(20)","NOT NULL, DEFAULT 'queued'","狀態機：queued/crawling/scanning/agent_testing/completed/failed/cancelled"),
            ("scan_mode","VARCHAR(20)","NOT NULL, DEFAULT 'passive'","掃描模式：passive/active"),
            ("max_depth","INTEGER","NOT NULL, DEFAULT 3","BFS 最大爬取深度"),
            ("max_pages","INTEGER","NOT NULL, DEFAULT 50","BFS 最大爬取頁數"),
            ("progress","JSONB","NOT NULL, DEFAULT '{}'","即時進度（pages_crawled/pages_scanned 等）"),
            ("overall_score","DECIMAL(5,2)","NULL","綜合評分（0–100）"),
            ("category_scores","JSONB","NULL","各維度評分 {SEO:xx, AEO:xx, GEO:xx, Security:xx}"),
            ("top_actions","JSONB","NULL","前 5 項改善建議清單"),
            ("created_at","TIMESTAMPTZ","NOT NULL","建立時間"),
            ("updated_at","TIMESTAMPTZ","NOT NULL","最後更新時間"),
         ]),
        ("表 8-2-7　scans_page 資料表",
         ["欄位名稱","資料型別","限制","說明"],
         [
            ("id","SERIAL","PK, NOT NULL","頁面唯一識別碼"),
            ("scan_id","UUID","FK→scans_scanjob, NOT NULL","關聯掃描任務"),
            ("url","VARCHAR(2048)","NOT NULL","頁面 URL"),
            ("status_code","INTEGER","NULL","HTTP 狀態碼"),
            ("title","VARCHAR(500)","NOT NULL, DEFAULT ''","頁面標題"),
            ("depth","INTEGER","NOT NULL, DEFAULT 0","BFS 爬取深度"),
            ("screenshot_path","VARCHAR(500)","NULL","截圖檔案路徑"),
            ("created_at","TIMESTAMPTZ","NOT NULL","爬取時間"),
         ]),
        ("表 8-2-8　scans_finding 資料表",
         ["欄位名稱","資料型別","限制","說明"],
         [
            ("id","SERIAL","PK, NOT NULL","發現事項唯一識別碼"),
            ("page_id","INTEGER","FK→scans_page, NOT NULL","關聯頁面"),
            ("category","VARCHAR(20)","NOT NULL","掃描維度：SEO/AEO/GEO/Security"),
            ("severity","VARCHAR(20)","NOT NULL","嚴重程度：high/medium/low/info"),
            ("title","VARCHAR(200)","NOT NULL","問題標題"),
            ("description","TEXT","NOT NULL","詳細說明"),
            ("recommendation","TEXT","NOT NULL, DEFAULT ''","改善建議"),
            ("created_at","TIMESTAMPTZ","NOT NULL","發現時間"),
         ]),
        ("表 8-2-9　reviews_platformreview 資料表",
         ["欄位名稱","資料型別","限制","說明"],
         [
            ("id","SERIAL","PK, NOT NULL","評論唯一識別碼"),
            ("user_id","INTEGER","FK→auth_user, UNIQUE（主評論）, NOT NULL","關聯使用者（每人一則主評論）"),
            ("rating","SMALLINT","NOT NULL, CHECK(1≤rating≤5)","評分（1–5 星）"),
            ("content","TEXT","NOT NULL","評論內容"),
            ("images","JSONB","NOT NULL, DEFAULT '[]'","附圖 URL 清單"),
            ("parent_id","INTEGER","FK→self, NULL","父評論（用於 thread 回覆結構）"),
            ("created_at","TIMESTAMPTZ","NOT NULL","評論時間"),
            ("updated_at","TIMESTAMPTZ","NOT NULL","最後修改時間"),
         ]),
        ("表 8-2-10　admin_api_adminauditlog 資料表",
         ["欄位名稱","資料型別","限制","說明"],
         [
            ("id","SERIAL","PK, NOT NULL","稽核日誌唯一識別碼"),
            ("actor_id","INTEGER","FK→auth_user, NOT NULL","執行操作的 Staff 管理員"),
            ("action","VARCHAR(100)","NOT NULL","操作類型（如：adjust_coins、reply_review）"),
            ("target_user_id","INTEGER","FK→auth_user, NULL","被操作的目標使用者"),
            ("detail","JSONB","NOT NULL, DEFAULT '{}'","操作詳細資料（金額、原因等）"),
            ("created_at","TIMESTAMPTZ","NOT NULL","操作時間（禁止刪除此表紀錄）"),
         ]),
    ]

    for caption, headers, rows in tables_meta:
        add_std_table(doc, headers, rows, caption)


# ── 參考資料 ────────────────────────────────────────

def _add_hyperlink(paragraph, url, text, color="0563C1", size=None):
    """於段落插入真正可點擊的 Word 超連結（外部 URL）"""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), CH_FONT)
    rFonts.set(qn('w:ascii'), EN_FONT)
    rFonts.set(qn('w:hAnsi'), EN_FONT)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int((size or BODY_SIZE).pt * 2)))
    rPr.append(sz)
    col = OxmlElement('w:color'); col.set(qn('w:val'), color); rPr.append(col)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    run.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _para_bottom_border(p, color="D0D0D0", sz=6):
    """於段落底部加分隔線（參考資料條目間區隔）"""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_references(doc):
    add_chapter(doc, "參考資料")
    add_body(doc, "以下參考資料以「[編號] 標題：連結」格式排列，連結皆為可點擊之真實來源，"
        "條目間以分隔線區隔；另標示人工智慧輔助使用情形。")

    # (標題, 可點擊連結)；查證日期 2026-05；無連結者為紙本專書
    refs = [
        ("Ahrefs — Pricing（Starter US$29/月，2026）", "https://ahrefs.com/pricing"),
        ("SEMrush — Pricing for SEO（Pro US$139.95/月，2026）", "https://www.semrush.com/pricing/seo/"),
        ("Moz — Moz Pro Pricing（Starter US$49/月，2026）", "https://moz.com/products/pro/pricing"),
        ("Sucuri — Website Security Platform Pricing（Basic US$199.99/年，2026）", "https://sucuri.net/website-security-platform/signup/"),
        ("Screaming Frog — SEO Spider Pricing（€245/年，2026）", "https://www.screamingfrog.co.uk/seo-spider/pricing/"),
        ("Google Search Console — 官方說明（免費）", "https://search.google.com/search-console/about"),
        ("經濟部中小及新創企業署 —《2024 年中小企業白皮書》（中小企業逾 167.4 萬家）", "https://www.sme.gov.tw/article-tw-2853-13097"),
        ("臺灣銀行 — 牌告匯率（外幣換算依據）", "https://rate.bot.com.tw/xrt"),
        ("OWASP Foundation — OWASP Top 10:2021", "https://owasp.org/Top10/"),
        ("Osterwalder, A., & Pigneur, Y. (2010). Business Model Generation. John Wiley & Sons.（紙本專書）", ""),
        ("Django Software Foundation — Django 5.1 documentation", "https://docs.djangoproject.com/en/5.1/"),
        ("Django REST Framework — documentation", "https://www.django-rest-framework.org/"),
        ("React — 官方文件", "https://react.dev/"),
        ("Vite — 官方文件", "https://vitejs.dev/"),
        ("Tailwind CSS — 官方文件", "https://tailwindcss.com/"),
        ("ReactFlow — 官方文件", "https://reactflow.dev/"),
        ("Microsoft — Playwright 官方文件", "https://playwright.dev/"),
        ("Celery Project — Celery 5.4 documentation", "https://docs.celeryq.dev/en/stable/"),
        ("Redis Ltd. — Redis 7.2 documentation", "https://redis.io/docs/"),
        ("PostgreSQL Global Development Group — PostgreSQL 16 documentation", "https://www.postgresql.org/docs/16/"),
        ("Docker Inc. — Docker documentation", "https://docs.docker.com/"),
        ("F5/NGINX — NGINX documentation", "https://nginx.org/en/docs/"),
        ("Astral — uv 官方文件（Python 套件管理）", "https://docs.astral.sh/uv/"),
        ("Git — 官方文件", "https://git-scm.com/doc"),
        ("python-docx — 官方文件", "https://python-docx.readthedocs.io/"),
        ("Matplotlib — 官方文件", "https://matplotlib.org/"),
        ("Google — Identity / OAuth 2.0 文件", "https://developers.google.com/identity"),
        ("PlantUML — 官方網站", "https://plantuml.com/"),
        ("OMG — Unified Modeling Language (UML) 2.5.1 規格", "https://www.omg.org/spec/UML/"),
    ]

    for i, (title, url) in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent       = Cm(1.2)
        p.paragraph_format.first_line_indent  = Cm(-1.2)
        run = p.add_run(f"[{i}] {title}")
        _set_run_font(run, BODY_SIZE)
        if url:
            sep = p.add_run("：")
            _set_run_font(sep, BODY_SIZE)
            _add_hyperlink(p, url, url)
        _set_para_spacing(p, before=2, after=4, ls=1.5)
        _para_bottom_border(p)

    doc.add_paragraph()
    add_body(doc, "※ 使用人工智慧輔助說明：", bold=True)
    add_std_table(doc,
        ["序號","使用工具名稱","使用範圍及說明","頁碼"],
        [
            ("1","Claude（Anthropic）","輔助系統手冊初稿內容撰寫、排版格式與圖表程式整理","全文"),
            ("2","Google Gemini","Hermes-Agent Phase 2 工具調用邏輯設計參考","第7章"),
        ],
        "表　人工智慧使用說明表")


# ── PlantUML 圖碼輸出 ────────────────────────────────

PLANTUML_CODE = """
================================================================
Argus 系統手冊 - PlantUML 圖碼集
使用方式：將各段 @startuml ... @enduml 複製至 https://www.plantuml.com/plantuml/uml/
================================================================

【圖 3-1-1】Argus 系統架構圖
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam componentStyle rectangle
title Argus 系統整體架構圖

package "用戶端層" {
  [瀏覽器 (Chrome/Firefox/Edge)]
}

package "反向代理層" {
  [Nginx 1.26\\n:80/:443]
}

package "應用程式層 (Django 5.1)" {
  [Django REST Framework\\nRESTful API]
  [React 18 SPA\\n(frontend/dist)]
  [BillingService\\n計費服務]
  [ReportGenerator\\nWord 報告生成]
}

package "非同步任務層" {
  [Celery Worker\\n掃描任務執行]
  [Redis 7.2\\n任務佇列]
  [Playwright Chromium\\nBFS 爬蟲]
  [四維掃描引擎\\nSEO/AEO/GEO/Security]
}

package "資料持久層" {
  database "PostgreSQL 16\\n10 個資料表"
}

[瀏覽器 (Chrome/Firefox/Edge)] --> [Nginx 1.26\\n:80/:443] : HTTPS
[Nginx 1.26\\n:80/:443] --> [Django REST Framework\\nRESTful API] : /api/*
[Nginx 1.26\\n:80/:443] --> [React 18 SPA\\n(frontend/dist)] : /*
[Django REST Framework\\nRESTful API] --> [BillingService\\n計費服務]
[Django REST Framework\\nRESTful API] --> [Redis 7.2\\n任務佇列] : 派發任務
[Redis 7.2\\n任務佇列] --> [Celery Worker\\n掃描任務執行]
[Celery Worker\\n掃描任務執行] --> [Playwright Chromium\\nBFS 爬蟲]
[Celery Worker\\n掃描任務執行] --> [四維掃描引擎\\nSEO/AEO/GEO/Security]
[Celery Worker\\n掃描任務執行] --> [ReportGenerator\\nWord 報告生成]
[BillingService\\n計費服務] --> "PostgreSQL 16\\n10 個資料表"
[Celery Worker\\n掃描任務執行] --> "PostgreSQL 16\\n10 個資料表"
@enduml

【圖 5-2-1】使用個案圖
@startuml
!theme plain
left to right direction
title Argus 使用個案圖

actor "一般使用者\\nUser" as User
actor "管理員 Staff" as Staff
actor "Celery Worker\\n（系統）" as Worker

rectangle Argus {
  usecase "登入系統\\n(Google OAuth)" as UC01
  usecase "提交掃描任務" as UC02
  usecase "查看掃描結果" as UC03
  usecase "下載健檢報告(.docx)" as UC04
  usecase "查看網站拓樸圖" as UC05
  usecase "購買點數" as UC06
  usecase "撰寫平台評論" as UC07
  usecase "管理使用者點數" as UC08
  usecase "回覆平台評論" as UC09
  usecase "管理 CMS 內容" as UC10
  usecase "執行 BFS 爬蟲" as UC11
  usecase "執行四維掃描" as UC12
}

User --> UC01
User --> UC02
User --> UC03
User --> UC04
User --> UC05
User --> UC06
User --> UC07
Staff --> UC08
Staff --> UC09
Staff --> UC10
Worker --> UC11
Worker --> UC12
UC02 .> UC11 : <<include>>
UC11 .> UC12 : <<include>>
@enduml

【圖 5-3-1】活動圖：提交掃描任務
@startuml
!theme plain
title 活動圖：提交掃描任務

|使用者|
start
:輸入目標 URL 與掃描設定;
|Django API|
:驗證 URL 格式;
if (URL 合法？) then (否)
  :回傳 HTTP 400;
  stop
endif
:檢查使用者點數;
if (點數充足？) then (否)
  :回傳 HTTP 402\\n引導購買;
  stop
endif
:hold_for_scan（預扣點數）;
:建立 ScanJob（status=queued）;
|Celery Worker|
:接收任務，status=crawling;
:Playwright BFS 爬取頁面;
:status=scanning;
:逐頁執行四維掃描引擎;
|Django API|
:settle_scan_actual（結算點數）;
:生成 Word 報告;
:status=completed;
|使用者|
:下載健檢報告;
stop
@enduml

【圖 5-4-1】分析類別圖
@startuml
!theme plain
title 分析類別圖（Analysis Class Diagram）

class User {
  +id: Integer
  +username: String
  +email: String
  +is_staff: Boolean
}

class ScanJob {
  +id: UUID
  +original_url: String
  +status: String
  +scan_mode: String
  +max_pages: Integer
  +overall_score: Decimal
}

class Page {
  +id: Integer
  +url: String
  +status_code: Integer
  +depth: Integer
}

class Finding {
  +id: Integer
  +category: String
  +severity: String
  +title: String
  +description: Text
  +recommendation: Text
}

class CoinWallet {
  +id: Integer
  +balance: Integer
}

class CoinTransaction {
  +id: Integer
  +amount: Integer
  +type: String
}

class PurchaseOrder {
  +id: UUID
  +status: String
}

class AgentSession {
  +id: Integer
  +status: String
}

User "1" -- "1" CoinWallet
User "1" -- "*" ScanJob
User "1" -- "*" PurchaseOrder
ScanJob "1" -- "*" Page
Page "1" -- "*" Finding
CoinWallet "1" -- "*" CoinTransaction
ScanJob ..> AgentSession : Phase 2
note on link: 虛線表示 Phase 2 功能
@enduml

【圖 6-1-1】掃描任務循序圖
@startuml
!theme plain
title 掃描任務循序圖（Sequential Diagram）

actor "User Browser" as User
participant "Django DRF\\n:8000" as DRF
participant "BillingService" as Billing
participant "Celery Queue\\n(Redis)" as Queue
participant "Celery Worker" as Worker
database "PostgreSQL" as DB

User -> DRF: POST /api/scans/\\n{url, scan_mode, max_pages}
DRF -> Billing: hold_for_scan(max_pages×10 coins)
Billing -> DB: SELECT FOR UPDATE wallet\\nUPDATE balance -= hold_amount
Billing --> DRF: ok
DRF -> DB: INSERT scans_scanjob (status=queued)
DRF -> Queue: apply_async(run_scan_task, scan_id)
DRF --> User: HTTP 201 {scan_id, status}

...非同步執行...

Queue -> Worker: 派發任務
Worker -> DB: UPDATE status=crawling
Worker -> Worker: Playwright BFS 爬蟲\\n抓取所有可達頁面
Worker -> DB: INSERT scans_page (批量)
Worker -> DB: UPDATE status=scanning
Worker -> Worker: 逐頁執行\\nSEO/AEO/GEO/Security 掃描
Worker -> DB: INSERT scans_finding (批量)
Worker -> Billing: settle_scan_actual(actual_pages×10)
Billing -> DB: UPDATE balance（退差額）
Worker -> Worker: 生成 Word 報告 (.docx)
Worker -> DB: UPDATE status=completed, overall_score

User -> DRF: GET /api/scans/{scan_id}/
DRF -> DB: SELECT scanjob + findings
DRF --> User: HTTP 200 {status, scores, findings}
@enduml

【圖 6-2-1】設計類別圖
@startuml
!theme plain
title 設計類別圖（Design Class Diagram）

class User {
  +id: int
  +username: str
  +email: str
  +is_staff: bool
  +date_joined: datetime
  --
  +get_wallet(): CoinWallet
}

class ScanJob {
  +id: UUID
  +user: FK[User]
  +original_url: str
  +status: str
  +scan_mode: str
  +max_depth: int
  +max_pages: int
  +progress: dict
  +overall_score: Decimal
  +category_scores: dict
  +top_actions: list
  --
  +cancel()
  +get_findings_summary(): dict
}

class BillingService <<service>> {
  --
  {static} +hold_for_scan(scan, max_pages)
  {static} +settle_scan_actual(scan, actual)
  {static} +refund_full_for_scan(scan)
  {static} +grant_monthly_bonus(user)
  {static} +purchase_coins(user, plan, order)
}

class CoinWallet {
  +id: int
  +user: FK[User]
  +balance: int
  +total_purchased_ntd: Decimal
  +last_bonus_year: int
  +last_bonus_month: int
  --
  +check_sufficient(amount): bool
}

class CoinTransaction {
  +id: int
  +wallet: FK[CoinWallet]
  +amount: int
  +type: str
  +scan: FK[ScanJob]
  +note: str
  +created_at: datetime
}

User "1" *-- "1" CoinWallet
User "1" *-- "*" ScanJob
CoinWallet "1" *-- "*" CoinTransaction
BillingService ..> CoinWallet : writes (only entry)
BillingService ..> CoinTransaction : writes (only entry)
ScanJob ..> BillingService : calls
@enduml

【圖 7-1-1】Docker Compose 佈署圖
@startuml
!theme plain
title Docker Compose 佈署圖（Deployment Diagram）

node "Client Browser" as browser

node "Cloud Server (AWS t3.xlarge)" {
  node "Docker Bridge Network" {
    component "nginx:1.26\\n:80/:443" as nginx
    component "web (Django 5.1)\\n:8000" as web
    component "worker (Celery 5.4)" as worker
    component "redis:7.2\\n:6379" as redis
    database "db (PostgreSQL 16)\\n:5432" as db
  }
}

browser --> nginx : HTTPS :443
nginx --> web : /api/* reverse proxy
nginx --> web : /* SPA fallback
web --> db : SQLAlchemy/Django ORM
web --> redis : 派發任務
worker --> redis : 接收任務
worker --> db : 寫入掃描結果
@enduml

【圖 7-2-1】套件架構圖
@startuml
!theme plain
title 套件架構圖（Package Diagram）

package "backend" {
  package "config" as cfg {
    [settings.py]
    [urls.py]
  }
  package "apps" {
    package "accounts" { [views.py\\nmodels.py] }
    package "scans" { [tasks.py\\ncrawler.py\\nscanners.py] }
    package "billing" { [services.py\\nmodels.py] }
    package "agent" { [loop.py\\nproviders.py] }
    package "reviews" { [models.py\\nviews.py] }
    package "admin_api" { [views.py\\npermissions.py] }
    package "content" { [models.py\\nadmin.py] }
  }
}

package "frontend/src" {
  [App.jsx]
  [api.js]
  [store.js]
}

cfg --> accounts
cfg --> scans
cfg --> billing
cfg --> agent
cfg --> reviews
cfg --> admin_api
cfg --> content
scans ..> billing : uses BillingService
agent ..> scans : Phase 2
@enduml

【圖 7-3-1】元件圖
@startuml
!theme plain
title 系統元件圖（Component Diagram）

component Browser {
  [React 18 SPA]
  [Axios API Client]
}
component "Nginx :443" as nginx
component "Django :8000" {
  [Django REST Framework]
  [BillingService]
  [ReportGenerator]
}
component "Celery Worker" {
  [ScanTaskRunner]
  [BFSCrawler]
  [ScanEngines\\nSEO/AEO/GEO/Security]
}
component "Redis :6379" as redis
database "PostgreSQL :5432" as pg
component "Playwright Chromium" as chromium #dashed

Browser --> nginx : HTTPS
nginx --> "Django :8000" : /api/* proxy
"Django REST Framework" --> BillingService
"Django REST Framework" --> redis : enqueue
redis --> "Celery Worker" : dequeue
BFSCrawler --> chromium : 爬取頁面
ScanTaskRunner --> pg : 讀寫結果
@enduml

【圖 7-4-1】ScanJob 狀態機圖
@startuml
!theme plain
title ScanJob 狀態機圖（State Machine）

[*] --> queued : ScanJob 建立

queued --> crawling : Worker 接收任務
crawling --> scanning : BFS 爬蟲完成
scanning --> agent_testing : Phase 2（選擇性）
scanning --> completed : 掃描完成（Phase 1）
agent_testing --> completed : Agent 測試完成

queued --> failed : 例外錯誤
crawling --> failed : 例外錯誤
scanning --> failed : 例外錯誤

queued --> cancelled : 使用者取消
crawling --> cancelled : 使用者取消
scanning --> cancelled : 使用者取消

completed --> [*]
failed --> [*]
cancelled --> [*]

note right of failed : refund_full_for_scan()\\n全退點數
note right of cancelled : refund_full_for_scan()\\n全退點數
@enduml

【圖 8-1-1】資料庫 ER 圖
@startuml
!theme plain
title 資料庫 ER 圖（Entity-Relationship Diagram）

entity auth_user {
  * id : SERIAL <<PK>>
  --
  username : VARCHAR(150) <<UK>>
  email : VARCHAR(254) <<UK>>
  is_staff : BOOLEAN
  is_active : BOOLEAN
  date_joined : TIMESTAMPTZ
}

entity billing_coinwallet {
  * id : SERIAL <<PK>>
  --
  user_id : INTEGER <<FK, UK>>
  balance : INTEGER
  total_purchased_ntd : DECIMAL
  last_bonus_year : INTEGER
  last_bonus_month : INTEGER
}

entity billing_cointransaction {
  * id : SERIAL <<PK>>
  --
  wallet_id : INTEGER <<FK>>
  scan_id : UUID <<FK, NULL>>
  amount : INTEGER
  type : VARCHAR(20)
  note : TEXT
  created_at : TIMESTAMPTZ
}

entity billing_pricingplan {
  * id : SERIAL <<PK>>
  --
  name : VARCHAR(100)
  coins : INTEGER
  price_ntd : DECIMAL
  is_active : BOOLEAN
}

entity billing_purchaseorder {
  * id : UUID <<PK>>
  --
  user_id : INTEGER <<FK>>
  plan_id : INTEGER <<FK>>
  transaction_id : INTEGER <<FK, NULL>>
  status : VARCHAR(20)
}

entity scans_scanjob {
  * id : UUID <<PK>>
  --
  user_id : INTEGER <<FK>>
  original_url : VARCHAR(2048)
  status : VARCHAR(20)
  scan_mode : VARCHAR(20)
  max_pages : INTEGER
  overall_score : DECIMAL
  category_scores : JSONB
  created_at : TIMESTAMPTZ
}

entity scans_page {
  * id : SERIAL <<PK>>
  --
  scan_id : UUID <<FK>>
  url : VARCHAR(2048)
  status_code : INTEGER
  depth : INTEGER
}

entity scans_finding {
  * id : SERIAL <<PK>>
  --
  page_id : INTEGER <<FK>>
  category : VARCHAR(20)
  severity : VARCHAR(20)
  title : VARCHAR(200)
  description : TEXT
}

entity reviews_platformreview {
  * id : SERIAL <<PK>>
  --
  user_id : INTEGER <<FK, UK>>
  rating : SMALLINT
  content : TEXT
  parent_id : INTEGER <<FK, NULL>>
}

entity admin_api_adminauditlog {
  * id : SERIAL <<PK>>
  --
  actor_id : INTEGER <<FK>>
  action : VARCHAR(100)
  target_user_id : INTEGER <<FK, NULL>>
  detail : JSONB
  created_at : TIMESTAMPTZ
}

auth_user ||--|| billing_coinwallet
auth_user ||--o{ scans_scanjob
auth_user ||--o{ billing_purchaseorder
auth_user ||--o{ reviews_platformreview
auth_user ||--o{ admin_api_adminauditlog
billing_coinwallet ||--o{ billing_cointransaction
scans_scanjob ||--o{ scans_page
scans_scanjob ||--o{ billing_cointransaction
scans_page ||--o{ scans_finding
billing_pricingplan ||--o{ billing_purchaseorder
reviews_platformreview |o--o{ reviews_platformreview : parent/child
@enduml
"""


# ── 主程式 ──────────────────────────────────────────

def main():
    doc = Document()
    set_page_margins(doc)
    add_page_number(doc)

    add_cover(doc)
    add_toc_pages(doc)

    ch1(doc); doc.add_page_break()
    ch2(doc); doc.add_page_break()
    ch3(doc); doc.add_page_break()
    ch4(doc); doc.add_page_break()
    ch5(doc); doc.add_page_break()
    ch6(doc); doc.add_page_break()
    ch7(doc); doc.add_page_break()
    ch8(doc); doc.add_page_break()
    add_references(doc)

    # 開啟文件時自動更新欄位（目錄/圖目錄/表目錄頁碼）
    settings = doc.settings.element
    upd = OxmlElement('w:updateFields'); upd.set(qn('w:val'), 'true')
    settings.append(upd)

    doc.save(OUT_DOCX)
    print(f"[OK] 已儲存：{OUT_DOCX}")

    with open(OUT_PLANTUML, "w", encoding="utf-8") as f:
        f.write(PLANTUML_CODE)
    print(f"[OK] PlantUML 圖碼：{OUT_PLANTUML}")


if __name__ == "__main__":
    main()
