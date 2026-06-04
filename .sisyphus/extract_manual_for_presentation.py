from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "專題文件" / "四技第 115401 組-ARGUS AI網站健檢平台-系統手冊.docx"
OUT_PATH = ROOT / ".sisyphus" / "manual_presentation_extract.json"


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u3000", " ")).strip()


def table_rows(table):
    rows = []
    for row in table.rows:
        rows.append([clean(cell.text) for cell in row.cells])
    return rows


def main() -> None:
    doc = Document(DOCX_PATH)
    paragraphs = []
    headings = []
    captions = []

    for i, p in enumerate(doc.paragraphs):
        text = clean(p.text)
        if not text:
            continue
        style = p.style.name if p.style is not None else ""
        item = {"index": i, "style": style, "text": text}
        paragraphs.append(item)
        if style.lower().startswith("heading") or re.match(r"^(第[一二三四五六七八九十]+章|\d+[-－]\d+|\d+\.\d+|[壹貳參肆伍陸柒捌玖拾])", text):
            headings.append(item)
        if re.match(r"^[▲▼]?\s*[圖表]\s*[\d\-－]+", text) or "圖 " in text or "表 " in text:
            captions.append(item)

    tables = []
    responsibility_tables = []
    for ti, table in enumerate(doc.tables):
        rows = table_rows(table)
        flat = " ".join(" ".join(r) for r in rows)
        item = {"index": ti, "rows": rows}
        tables.append(item)
        if any(k in flat for k in ["分工", "組員", "姓名", "負責", "工作項目", "貢獻"]):
            responsibility_tables.append(item)

    rels = doc.part.rels
    image_count = sum(1 for rel in rels.values() if "image" in rel.reltype)

    # Keep body excerpts compact enough for the model to inspect safely.
    chapter_markers = []
    for item in paragraphs:
        if re.search(r"(第一章|第二章|第三章|第四章|第五章|第六章|第七章|第八章|第九章|第十章)", item["text"]):
            chapter_markers.append(item)

    data = {
        "docx": str(DOCX_PATH),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "image_count": image_count,
        "headings": headings[:300],
        "chapter_markers": chapter_markers[:80],
        "captions": captions[:200],
        "responsibility_tables": responsibility_tables,
        "tables_preview": tables[:30],
        "paragraphs_preview": paragraphs[:500],
    }
    OUT_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({
        "docx": str(DOCX_PATH),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "image_count": image_count,
        "responsibility_table_count": len(responsibility_tables),
        "out": str(OUT_PATH),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
